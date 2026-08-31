from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "GradTrack_Test_Plan.docx"

PAGE_WIDTH_TWIPS = 12240
PAGE_HEIGHT_TWIPS = 15840
USABLE_WIDTH_TWIPS = 8640  # Letter width less 1.5-inch left and 1-inch right margins.
SYSTEM_WIDTHS = (1440, 1584, 1512, 2952, 1152)
BLACK_BOX_WIDTHS = (1152, 2232, 3960, 1296)


def bullets(value: str | Sequence[str]) -> list[str]:
    if isinstance(value, str):
        return [part.strip() for part in value.split("|") if part.strip()]
    return [str(part).strip() for part in value if str(part).strip()]


DB = "Database connection must be established."
SERVER = "Frontend, PHP API, and required server services must be available."
ADMIN_ACCOUNT = "An active Admin-role account must exist."
SUPER_ACCOUNT = "An active Super Admin account must exist."
ALUMNI_ADMIN_ACCOUNT = "An active Alumni Admin account must exist."
REGISTRAR_ACCOUNT = "An active Registrar account must exist."
DEAN_ACCOUNT = "An active Dean-CCS, Dean-COED, or Dean-HM account must exist."
GRAD_ACCOUNT = "An approved and active Graduate account must exist."


web_cases: list[tuple[str, list[str], list[str]]] = []


def web(description: str, prereq: str | Sequence[str], expected: str | Sequence[str]) -> None:
    web_cases.append((description, bullets(prereq), bullets(expected)))


# Public website and authentication
web("Public Website - Home page and primary navigation", [DB, SERVER],
    "The configured GradTrack name, branding, and public content must load.|Desktop navigation links must open the correct public pages.|Survey and sign-in calls to action must use valid routes.|A server failure must show a usable error or fallback without exposing technical details.")
web("Public Website - Published announcement list", [DB, "Published public announcements must exist."],
    "Only published public announcements must be listed.|Search, category filtering, pagination, and clear-filter behavior must work.|An empty search must display a clear no-results state.|Draft and archived announcements must not be exposed.")
web("Public Website - Announcement details", [DB, "A published public announcement must exist."],
    "The selected announcement title, summary, content, date, category, and available media must load.|An unknown announcement ID must show an unavailable/not-found state.|Recent announcement links must open the correct record.")
web("Public Website - About page", [DB, "About-page content must be configured in System Settings."],
    "Visible About sections must appear in the configured order.|Configured image and call-to-action content must render correctly.|Hidden or empty optional sections must not create broken blocks.")
web("Public Website - Frequently Asked Questions page", [DB, "FAQ categories and items must be configured."],
    "Visible FAQ categories and questions must load in their configured order.|Expanding and collapsing a question must work.|Hidden FAQ items must not be displayed.|An empty FAQ configuration must produce a clean empty state.")
web("Public Website - Privacy Policy page", [DB, "Privacy Policy content must be configured."],
    "The introduction, effective date, updated date, and visible sections must load.|Allowed formatted content must render without unsafe script execution.|Hidden sections must not be displayed.")
web("Admin Authentication - Sign in", [DB, "An active administrator account must exist."],
    "The email and password controls, password visibility control, and forgot-password link must be available.|Valid credentials must authenticate the account and redirect it according to its actual role.|Invalid credentials must display a generic error.|Inactive accounts must be denied.")
web("Graduate Authentication - Sign in", [DB, GRAD_ACCOUNT],
    "The Graduate sign-in interface must load with email, password, show/hide, and forgot-password controls.|Approved active accounts must be redirected to Graduate Announcements.|Pending, rejected, and inactive accounts must be denied with the appropriate status message.|Invalid credentials must not disclose account secrets.")
web("Admin Authentication - Forgot password OTP request", [DB, "A registered administrator email must exist."],
    "The request form must accept a valid email address.|A generic response must be used for registered and unregistered addresses.|A six-digit OTP must be sent only when the account is eligible.|Resend must observe the one-minute cooldown.")
web("Admin Authentication - OTP verification and password reset", [DB, "A valid administrator password-reset OTP must have been requested."],
    "A valid six-digit OTP must advance to password reset.|Expired OTPs and more than five invalid attempts must be rejected.|The new password must require at least eight characters with uppercase, lowercase, number, and symbol.|Mismatched passwords must be rejected and a successful reset must return to Admin Sign In.")
web("Graduate Authentication - Forgot password OTP request", [DB, "A registered Graduate email must exist."],
    "The request form must validate the email format.|A generic response must prevent account enumeration.|A six-digit OTP must be sent only to an eligible Graduate account.|Resend must be unavailable during the cooldown.")
web("Graduate Authentication - OTP verification and password reset", [DB, "A valid Graduate password-reset OTP must have been requested."],
    "The valid OTP must open the reset step.|Expired, malformed, or repeatedly incorrect OTPs must be rejected.|Password complexity and confirmation matching must be enforced.|A successful reset must allow sign-in with the new password and invalidate the old password.")
web("Authentication - Role-based landing pages", [DB, "Active accounts for each implemented administrator role must exist."],
    "Admin must reach its dashboard.|Registrar must reach Manage Graduates.|Alumni Admin must reach Alumni Verification.|Super Admin must reach User Management.|Dean roles must reach Survey Participation.|MIS Staff and Research Coordinator must reach their available dashboard.")
web("Authentication - Protected route enforcement", [DB, "At least two accounts with different roles must exist."],
    "Unauthenticated requests must redirect to the correct sign-in page.|Authenticated users must be denied routes outside their allowed role.|The system must use a safe role-specific fallback page.|Direct URL entry must not bypass backend authorization.")
web("Authentication - Session persistence and logout", [DB, "A user must be authenticated."],
    "A valid session must persist during normal navigation and page refresh.|Logout confirmation must allow canceling without ending the session.|Confirmed logout must invalidate local authentication data and return to the correct sign-in page.|Protected pages must remain inaccessible after logout.")
web("Authentication - Maintenance-mode access", [DB, "Maintenance mode must be configurable by Super Admin."],
    "Public protected functions and Graduate Portal access must show the maintenance page while enabled.|Non-Super-Admin administrator roles must be blocked.|Super Admin must retain access to restore the system.|Normal access must resume after maintenance mode is disabled.")

# Tracer survey and optional account creation
web("Tracer Survey - Active survey verification interface", [DB, "An active survey and matching Registrar graduate record must exist."],
    "The active survey and identity-verification form must load.|The user may verify with student number in 2XXX-XXXX format or registered email, plus last name and program.|Matching details must issue a time-limited survey token.|Mismatched or missing records must be rejected.")
web("Tracer Survey - Verification status restrictions", [DB, "Submitted, expired-token, and unmatched test records must exist."],
    "A graduate who already submitted the selected survey must not obtain a duplicate token.|No active survey must produce a clear unavailable state.|An invalid or expired token must return the user to verification.|A token must not grant access to a different survey or graduate.")
web("Tracer Survey - Consent and section navigation", [DB, "A valid unused survey token must exist."],
    "The data-privacy agreement must be accepted before the questionnaire starts.|Question sections must appear in their configured order.|Previous and Next navigation must preserve entered answers.|The progress indicator must correspond to the current section.")
web("Tracer Survey - Supported question types", [DB, "The active survey must contain text, date, choice, checkbox, radio, rating, and header items."],
    "Each configured question type must render the correct control.|Required and optional indicators must match survey configuration.|Choice options and rating values must be selectable.|Header items must display as non-answer section text.")
web("Tracer Survey - Conditional employment questions", [DB, "A valid unused survey token must exist."],
    "Employment, training, and graduate-study branches must appear only when their controlling answers require them.|Hidden branch questions must not block navigation.|Changing a controlling answer must clear or ignore answers that are no longer applicable.")
web("Tracer Survey - Philippine address selection", [DB, "PSGC location data must be available."],
    "Region selection must load its provinces.|Province selection must load the correct cities/municipalities.|City/municipality selection must load the correct barangays.|Required location fields must reject incomplete or inconsistent selections.")
web("Tracer Survey - Required and Other-option validation", [DB, "The active survey must include required fields and an Other choice."],
    "A section with unanswered required questions must not advance.|Selecting Other must require its supporting text.|Removing Other must remove the supporting-text requirement.|Validation messages must identify the fields requiring correction.")
web("Tracer Survey - Draft restoration", [DB, "A valid survey token and partially completed browser draft must exist."],
    "Entered responses must be saved locally while progressing.|Refreshing or reopening the valid token must offer/restore the available draft.|Restored data must map to the correct questions and graduate.|A submitted or invalid token must not restore an editable draft.")
web("Tracer Survey - Successful submission", [DB, "A fully completed valid survey response must exist in the browser."],
    "The final review/submit action must validate all required answers.|The response and normalized location data must be stored once.|The survey token must be marked submitted.|A success view must be shown and repeat submission must be blocked.")
web("Tracer Survey - Server and duplicate-submission handling", [DB, "A valid survey token must exist."],
    "An API failure must preserve entered data and show a recoverable error.|A duplicate or already-used token must not create another response.|Malformed answers must be rejected by the server.|No partial response must be reported as successfully submitted.")
web("Graduate Account Registration - Form after survey", [DB, "A survey response without a linked Graduate account must have been submitted."],
    "The optional registration form must use the submitted graduate identity.|A valid unique email and complex matching password must create one account.|The new account must have pending alumni-verification status.|The user must be told that approval is required before portal access.")
web("Graduate Account Registration - Validation and duplicate prevention", [DB, "A submitted survey response must exist."],
    "Invalid email, weak password, and password mismatch must be rejected.|An email or graduate already linked to an account must not create a duplicate.|A response ID that does not belong to the graduate must be rejected.|API errors must not create a partially linked account.")

# Graduate portal
web("Graduate Portal - Announcements list", [DB, GRAD_ACCOUNT, "Published Alumni Admin announcements must exist."],
    "Published Graduate announcements must load after sign-in.|Search, category, pagination, and clear-filter actions must work.|Draft and archived announcements must remain hidden.|No-result and server-error states must be understandable.")
web("Graduate Portal - Announcement details and gallery", [DB, GRAD_ACCOUNT, "A published announcement with gallery media must exist."],
    "The selected announcement and recent announcements must load.|Cover and gallery images must use the configured fallback when missing.|Image preview must support next, previous, and close.|An unavailable record must not crash the portal.")
web("Graduate Portal - Dashboard summary", [DB, GRAD_ACCOUNT, "Graduate profile and survey-derived data must exist."],
    "Alumni score, employment/alignment indicators, badges, and information tiles must load from the Graduate's data.|Tiles must navigate only to enabled portal modules.|Missing optional profile data must be shown without a rendering failure.")
web("Graduate Profile - View own profile", [DB, GRAD_ACCOUNT],
    "The Graduate must see basic details and read-only survey-derived employment, education, and training information.|Missing optional values must display cleanly.|Editing controls must be available only on the owner's profile.")
web("Graduate Profile - Edit basic information", [DB, GRAD_ACCOUNT],
    "Valid name, email, phone, and address changes must be saved.|Email format and duplicate email must be validated.|A failed update must not overwrite the last saved profile.|The refreshed profile must show the stored values.")
web("Graduate Profile - Public alumni profile", [DB, GRAD_ACCOUNT, "Another Graduate profile must exist."],
    "The selected Graduate's public profile information must load.|Private editing and password controls must not be exposed.|Invalid Graduate IDs must show an unavailable profile state.|The viewer may start a direct chat only when messaging is enabled.")
web("Graduate Profile - Profile-picture upload and removal", [DB, GRAD_ACCOUNT],
    "JPG, PNG, WebP, or GIF images up to 5 MB must upload and display.|Invalid file type and oversized file must be rejected.|Removing the image must require the intended action and restore the default avatar.|A server failure must retain the previous image.")
web("Graduate Profile - Cover-photo upload and removal", [DB, GRAD_ACCOUNT],
    "A valid supported cover image up to 5 MB must upload and display.|Invalid or oversized media must be rejected before replacement.|Removal must restore the default cover presentation.|The cover change must be visible after refresh.")
web("Graduate Profile - Change password", [DB, GRAD_ACCOUNT],
    "The correct current password and a complex matching new password must update the credential.|Incorrect current password, weak password, and mismatch must be rejected.|The old password must fail after success and the new password must authenticate.")
web("Community Forum - Feed, search, and filters", [DB, GRAD_ACCOUNT, "Community Forum must be enabled and approved posts must exist."],
    "Approved posts must load in the feed.|Search and category, program, and graduation-year filters must narrow results correctly.|Clearing filters must restore the list.|No-result and loading states must be displayed.")
web("Community Forum - Create text post", [DB, GRAD_ACCOUNT, "Community Forum must be enabled."],
    "A non-empty title, content, and category must create a post when content passes moderation.|Required-field and prohibited-content validation must be shown.|The created post must be associated with the authenticated Graduate.|A failed request must not create a duplicate post.")
web("Community Forum - Create post with media", [DB, GRAD_ACCOUNT, "Forum media uploads must be enabled."],
    "Up to ten supported images or videos must be accepted.|Images over 5 MB, videos over 50 MB, invalid types, and excess files must be rejected.|Upload preview and removal before submission must work.|Accepted media must appear with the created post.")
web("Community Forum - Edit and delete own post", [DB, GRAD_ACCOUNT, "The Graduate must own an existing forum post."],
    "Editing must load the existing values and validate changes.|An edited post must be returned to pending moderation as implemented.|Cancel must leave the record unchanged.|Confirmed delete must remove the post and related visible content, while another user's post must not be editable or deletable.")
web("Community Forum - Likes and comments", [DB, GRAD_ACCOUNT, "An approved post must exist."],
    "Like and unlike must update the count once per Graduate.|A non-empty comment must appear under the correct post.|Blank comments must be rejected.|A Graduate may delete their own comment with confirmation but not another Graduate's comment.")
web("Community Forum - Post and comment reporting", [DB, GRAD_ACCOUNT, "A reportable post and comment must exist."],
    "The report dialog must identify the selected post or comment.|A reason and optional details must be submitted to moderation.|Cancel must not create a report.|Repeated or failed submissions must not falsely report success.")
web("Community Forum - Media viewer", [DB, GRAD_ACCOUNT, "A forum post with multiple media items must exist."],
    "Selecting media must open the correct image/video preview.|Image zoom and next/previous navigation must work where available.|Closing the viewer must return to the same post position.|Missing media must use a safe error state.")
web("Messaging - Alumni directory and direct conversation", [DB, GRAD_ACCOUNT, "Community Forum and Messaging must be enabled."],
    "The directory must support name, program, and batch search/filtering.|Selecting one Graduate must create or open the existing direct room.|Duplicate direct rooms for the same pair must not be created.|The current Graduate must not be offered as a new recipient.")
web("Messaging - Group-chat creation", [DB, GRAD_ACCOUNT, "At least two other approved Graduate accounts must exist."],
    "A group name and at least two selected members must create a group room.|Empty name or too few members must be rejected.|Selected members must be shown before creation.|Cancel must not create a room.")
web("Messaging - Conversation list and unread state", [DB, GRAD_ACCOUNT, "Direct and group rooms with messages must exist."],
    "Rooms must show latest-message information and unread counts.|Search must find the intended conversation.|Opening a room must load the correct participants and messages.|Read state must update without affecting another room.")
web("Messaging - Send and receive text in real time", [DB, GRAD_ACCOUNT, "Socket.IO messaging service must be connected."],
    "A non-empty text message up to 5,000 characters must be sent to the active room.|The sender must see pending then confirmed status.|Other room members must receive the message without refreshing.|Blank or over-limit text must not be accepted.")
web("Messaging - Image attachment messages", [DB, GRAD_ACCOUNT, "An active chat room and Socket.IO connection must exist."],
    "A supported image up to 10 MB must upload and appear as its own message item.|Image and text content must not be incorrectly merged.|Selecting an image must open its preview and Close must dismiss it.|Invalid or oversized images must be rejected.")
web("Messaging - Document attachments", [DB, GRAD_ACCOUNT, "An active chat room must exist."],
    "Supported PDF, Office, text, and CSV files up to 25 MB must upload with the original name.|Dangerous extensions and mismatched MIME types must be rejected.|The recipient must receive a usable attachment link.|A failed upload must show retry/failure without a false sent state.")
web("Messaging - Typing indicator", [DB, GRAD_ACCOUNT, "Two Graduate sessions must be in the same room with Socket.IO connected."],
    "Typing must be announced to other room participants while input is active.|The indicator must stop after inactivity, send, room change, or disconnect.|A user must not see their own typing indicator.")
web("Messaging - Presence and last seen", [DB, GRAD_ACCOUNT, "At least two Graduate sessions must be available."],
    "Connected users must show online status.|Disconnected users must transition to offline and show the stored last-seen time.|Away/offline state must remain consistent after navigation and reconnect.|Presence from another account must not overwrite the current user.")
web("Messaging - Delivery, read receipts, and synchronization", [DB, GRAD_ACCOUNT, "A room with two active users must exist."],
    "Message delivery and read status must update for the correct message.|Older messages must load without reordering or duplication.|A new-message indicator must work while reading older content.|Refresh must reproduce the server message history.")
web("Messaging - Reconnection and failed-message recovery", [DB, GRAD_ACCOUNT, "A chat room must be open and the realtime connection must be interruptible."],
    "Connection loss must be shown without deleting composed text.|A failed optimistic message must be marked and offer retry.|After reconnection, missed messages must synchronize once and in order.|Unread counts and presence must recover consistently.")
web("Alumni Job Support - Browse approved jobs", [DB, GRAD_ACCOUNT, "Approved active job posts must exist and Job Support must be enabled."],
    "Only approved and active job posts must be visible.|Search and title, company, skills, location, and program-fit filters must work.|No-result state and filter clearing must work.|Pending, declined, or inactive posts must not appear.")
web("Alumni Job Support - Job details and external application information", [DB, GRAD_ACCOUNT, "An approved job with application information must exist."],
    "The detail view must show company, type, location, industry, salary, skills, program fit, dates, and available poster information.|At least one configured application method, link, email, or requirements file must be shown.|Expired/missing jobs must not open as valid opportunities.")
web("Alumni Job Support - Job-posting eligibility", [DB, GRAD_ACCOUNT, "Job Support must be enabled."],
    "An employed Graduate with the required alumni-score permission must access Job Posting.|A Graduate without permission must see the locked explanation.|Direct URLs and APIs must not bypass the posting permission.")
web("Alumni Job Support - Create job post", [DB, GRAD_ACCOUNT, "The Graduate must have job-posting permission."],
    "Title, company, description, and at least one application contact/method must be required.|Application email must use a valid format.|Optional type, location, industry, salary, skills, program fit, deadline, and active status must save.|A valid post must enter pending approval and remain hidden from Browse Jobs.")
web("Alumni Job Support - Edit and delete own job post", [DB, GRAD_ACCOUNT, "The Graduate must own an existing job post."],
    "Editing must prefill current data and resubmit the item for approval.|Approval notes/status must be visible to the owner.|Cancel must preserve data.|Confirmed delete must remove the owner's post, while another Graduate's post must be protected.")
web("Graduate Notifications - Bell, list, and read state", [DB, GRAD_ACCOUNT, "Notifications must be enabled and notification records must exist."],
    "The bell must display the correct unread count.|Opening the list must show relevant announcement, survey, approval, forum, and job events.|Mark-one and mark-all-read must persist.|Selecting a notification must navigate to a valid permitted destination.")
web("Graduate Portal - Feature-flag behavior", [DB, GRAD_ACCOUNT, "Feature flags must be editable by Super Admin."],
    "Disabled Forum, Messaging, Job Support, Notifications, or survey features must be hidden or show the defined unavailable state.|Dependent Messaging must remain unavailable when Community Forum is disabled.|Re-enabling a feature must restore its navigation without granting extra role permissions.")

# Administrator common and Admin-role workflows
web("Administrator Profile - View and edit", [DB, "An administrator of any implemented role must be authenticated."],
    "The profile must display the authenticated administrator's role and email.|A valid full-name change must save while email remains read-only.|Profile-image upload must accept JPG, PNG, WebP, or GIF up to 5 MB and reject other/oversized files.|Server errors must leave saved data unchanged.")
web("Administrator Profile - Change password", [DB, "An administrator of any implemented role must be authenticated."],
    "Correct current password and a complex matching new password must update the credential.|Incorrect current password, weak password, or mismatch must be rejected.|The old password must no longer authenticate after success.")
web("Administrator Layout - Navigation, theme, notifications, and logout", [DB, "An administrator must be authenticated."],
    "The sidebar must show only links allowed for the current role.|Theme changes must update the interface consistently.|Notification count/read behavior must work when enabled.|Confirmed logout must end the session and Cancel must keep it active.")
web("Admin Dashboard - Tracer and employment summary", [DB, ADMIN_ACCOUNT, "Survey and graduate response data must exist."],
    "The selected-survey snapshot, total figures, employability index by program, employment trend, and job-alignment distribution must load.|Survey/year/program selections must update the same data scope.|Empty data and API failures must not display invented figures.")
web("MIS Staff and Research Coordinator - Available dashboard", [DB, "An active MIS Staff or Research Coordinator account must exist."],
    "The implemented dashboard must load for the role.|Survey-management, report, graduate CRUD, and Super Admin links must not be shown unless separately authorized.|Direct access to unauthorized administrator APIs must be denied.")
web("Admin Graduate Participation - Search and filters", [DB, ADMIN_ACCOUNT, "A survey with answered and non-respondent graduates must exist."],
    "Survey, response-status, program, graduation-year, and name/email/student-ID filters must return matching participants.|Pagination must preserve filters.|A no-result search must show an empty state.|Clear/reset must restore the default non-respondent view.")
web("Admin Graduate Participation - Selection and email notification", [DB, ADMIN_ACCOUNT, "Eligible non-respondent graduates with email addresses must exist."],
    "Visible rows and all matching rows must be selectable according to the current filter.|Notify must require at least one eligible selection and confirmation.|The system must report sent, skipped, and failed outcomes.|Answered or missing-email records must not be falsely reported as notified.")
web("Admin Graduate Participation - Submitted-answer view", [DB, ADMIN_ACCOUNT, "At least one completed survey response must exist."],
    "View Answers must open the selected Graduate's correct survey responses and identity summary.|Non-respondents must not have a fabricated answer view.|Closing the modal must return to the filtered list.")
web("Survey Management - Survey list and single-active rule", [DB, ADMIN_ACCOUNT, "Survey records in draft, inactive, and active states must exist."],
    "Survey cards and statuses must load.|Only one survey may be active at a time.|Creating or activating another survey while one is active must be blocked with a clear instruction.|Draft and inactive surveys must remain editable.")
web("Survey Management - Create survey from tracer template", [DB, ADMIN_ACCOUNT, "No conflicting active survey may exist."],
    "A title, description, and valid status must create a survey using the implemented tracer-question template.|Required metadata must be validated.|Cancel must not create a survey.|A failed request must not leave an incomplete survey.")
web("Survey Management - Edit sections and questions", [DB, ADMIN_ACCOUNT, "An editable survey must exist."],
    "The Admin must add, edit, remove, and order sections/questions.|Text, date, multiple-choice, radio, rating, checkbox, and header types must retain their configuration.|Choice questions must retain options and required flags.|Preview must match the saved questionnaire.")
web("Survey Management - Delete survey and clear surveys", [DB, ADMIN_ACCOUNT, "Deletable survey records must exist."],
    "Delete must identify that linked tokens/responses are affected and require confirmation.|Cancel must preserve the survey.|Confirmed delete must remove only the selected survey and dependencies.|Clear All Surveys must require its stronger confirmation and report errors safely.")
web("Survey Responses - List, details, and CSV export", [DB, ADMIN_ACCOUNT, "Completed responses must exist for a survey."],
    "The selected survey's responses and Graduate identity details must load.|Search/filter and response detail view must show correct stored answers.|CSV export must contain the intended response set and safe cell values.|No-response and export-error states must be handled.")
web("Survey Analytics - Charts and question analysis", [DB, ADMIN_ACCOUNT, "A survey with responses must exist."],
    "Response totals, rates, program/year/employment summaries, and question-level analysis must match stored responses.|Changing the survey must refresh all analytics.|Unsupported/empty answers must not break charts.|Excel/PDF analytics exports must represent the selected survey.")
web("Reports and Analytics - Overview filters", [DB, ADMIN_ACCOUNT, "Graduate, employment, and survey data must exist."],
    "Employment, alignment, graduation year, program, region, province, city/municipality, and barangay filters must apply together.|Dependent locations must update correctly.|Reset must restore the unfiltered report.|Summary values and tables must use the same result set.")
web("Reports and Analytics - Report tabs", [DB, ADMIN_ACCOUNT, "Reportable data must exist."],
    "Overview, By Program, By Year, Employment Status, Salary Distribution, Location, and Survey Analytics tabs must load.|Tab-specific year/department controls must update data.|Charts, totals, and tables must remain consistent.|Empty states must not invent percentages.")
web("Reports and Analytics - Excel and PDF export", [DB, ADMIN_ACCOUNT, "A report and filters must be selected."],
    "Excel and PDF exports must download successfully.|Exported content must include the selected scope, summary, tables/charts where supported, and filter metadata.|Spreadsheet text must be handled safely.|An export failure must show an error without a corrupt success file.")
web("Reports and Analytics - AI descriptive analysis", [DB, ADMIN_ACCOUNT, "Report data and the configured AI service must be available."],
    "Generated analysis must use the selected report data and identify itself as descriptive support.|Changing filters must not reuse an unrelated prior analysis.|Service timeout or configuration failure must show an unavailable/retry state without blocking standard reports.")

# Registrar and Dean workflows
web("Registrar - Graduate list, program tabs, search, and year filter", [DB, REGISTRAR_ACCOUNT, "Graduate records and program data must exist."],
    "Program tabs must show only matching records.|Name, student ID, or email search and graduation-year filtering must work with pagination.|Clear filters must restore the program list.|No-match results must show a clean empty state.")
web("Registrar - Add graduate record", [DB, REGISTRAR_ACCOUNT],
    "First name, last name, program, graduation year, and phone must be required.|Student ID must match four digits-hyphen-four digits and phone must be 11 digits beginning 09.|Optional email must be valid and unique when entered.|A valid record must save once and appear under the correct program/year.")
web("Registrar - Edit graduate record", [DB, REGISTRAR_ACCOUNT, "An existing Graduate record must be present."],
    "The edit form must load the selected record.|Valid permitted changes must save and retain related identity correctly.|Duplicate student ID/email or invalid phone/year must be rejected.|Cancel or failed save must keep the prior record.")
web("Registrar - Delete one or multiple graduate records", [DB, REGISTRAR_ACCOUNT, "Deletable Graduate records must exist."],
    "Single and selected-record deletion must require confirmation.|Cancel must preserve all records.|Confirmed deletion must remove only the selected records and return a correct count.|Related-data restrictions or server errors must be reported without partial false success.")
web("Registrar - Delete graduates by program and year", [DB, REGISTRAR_ACCOUNT, "Records must exist for a selected program/year."],
    "The destructive scope must be displayed before confirmation.|Cancel must preserve the cohort.|Confirmed deletion must affect only the chosen program and graduation year.|An empty cohort must not report records deleted.")
web("Registrar - Import graduates from Excel", [DB, REGISTRAR_ACCOUNT, "A readable XLS or XLSX file must be available."],
    "The first worksheet and supported header aliases must be parsed.|Valid names, student IDs, contact details, program, year, and employment values must be mapped.|Rows missing required name/year, invalid rows, and duplicate student ID/email must be reported/skipped according to implementation.|The result summary must distinguish imported, skipped, and failed rows.")
web("Dean Survey Participation - Program-scoped list", [DB, DEAN_ACCOUNT, "Survey participation records must exist across several programs."],
    "Dean-CCS must be limited to BSCS and ACT, Dean-COED to BSED and BEED, and Dean-HM to BSHM.|Search, response-status, year, and pagination controls must work inside the allowed programs.|Direct API filters must not expose other programs.")
web("Dean Survey Participation - Notify and view answers", [DB, DEAN_ACCOUNT, "Allowed-program respondents and non-respondents must exist."],
    "The Dean may select and notify eligible non-respondents in permitted programs.|Completed responses in permitted programs may be viewed.|Unpermitted programs and fabricated answers must remain inaccessible.|Send results must distinguish sent, failed, and skipped recipients.")

# Super Admin workflows
web("User Management - Search and filters", [DB, SUPER_ACCOUNT, "Administrator accounts of multiple roles and statuses must exist."],
    "Search by name, username, or email and role/status filters must return matching administrator accounts.|Pagination must preserve filters.|No-result and clear-filter behavior must work.|Graduate accounts must not appear in this administrator table.")
web("User Management - Create administrator", [DB, SUPER_ACCOUNT],
    "Username, email, password, full name, role, and status must be validated.|Unique username/email and a valid implemented role must be required.|A valid account must save once with the selected active state.|Weak/invalid data or an API failure must not create a partial user.")
web("User Management - Edit role and account information", [DB, SUPER_ACCOUNT, "An administrator account must exist."],
    "The selected user's data must prefill the form.|Valid username, email, name, role, active status, and optional new password changes must save.|Duplicate identifiers and invalid role/password must be rejected.|Cancel must preserve the stored account.")
web("User Management - Activate and deactivate", [DB, SUPER_ACCOUNT, "An active and an inactive administrator account must exist."],
    "Status changes must require confirmation.|Deactivation must prevent future authentication and activation must restore it.|The logged-in Super Admin must not be allowed to deactivate their own account.|No delete control must be presented by the current interface.")
web("Auto Email Reminders - Status and eligible-graduate list", [DB, SUPER_ACCOUNT, "An active survey with respondent and non-respondent records must exist."],
    "The active survey, eligible count, configured interval, and sent/failed/skipped statistics must load.|The eligible list must contain only graduates who have not answered and can be contacted.|No active survey must show a clear instruction.")
web("Auto Email Reminders - Manual send", [DB, SUPER_ACCOUNT, "Email delivery must be enabled and eligible non-respondents must exist."],
    "Subject and message may use the provided defaults or entered text.|Confirmation must identify the active survey and eligible count.|The result must report eligible, sent, failed, and skipped totals.|No active survey, disabled email, or zero eligible recipients must prevent sending.")
web("Auto Email Reminders - Frequency and logs", [DB, SUPER_ACCOUNT],
    "Preset or custom intervals from 1 through 365 days must save.|Out-of-range input must be constrained/rejected.|History must show auto/manual type, recipient, status, and date for logged attempts.|The interface must state that scheduled execution requires the configured server task and secret.")
web("Audit Trail - Search, filters, and details", [DB, SUPER_ACCOUNT, "Audit records must exist."],
    "Search and role, department, action, module, and start/end date filters must return matching logs.|Pagination, refresh, clear filters, and empty state must work.|The detail view must show available previous values, new values, and metadata without exposing secrets.")
web("Audit Trail - CSV export", [DB, SUPER_ACCOUNT, "A filtered audit result must exist."],
    "CSV export must use the current filters and include the intended audit fields.|Values beginning with spreadsheet formulas must be neutralized.|No-record and server-failure cases must not produce a misleading successful export.")
web("Database Backup - Summary and SQL download", [DB, SUPER_ACCOUNT],
    "Database name/summary, table count, record count, size, and table list must load.|Refresh must request current values.|Backup must download an SQL file containing supported structure and data.|Failure must show an error and unauthorized roles must be denied.")
web("System Settings - General configuration", [DB, SUPER_ACCOUNT],
    "System name, short name, institution, login title, description, contact details, footer, and copyright values must load and save.|Required names/title and email format must be enforced.|Cancel/failed save must not replace the active configuration.|Updated public branding text must appear after refresh.")
web("System Settings - Branding assets and colors", [DB, SUPER_ACCOUNT],
    "Valid logos/backgrounds must accept PNG, JPG, WebP, or GIF and favicon may also accept ICO.|Background must be at most 5 MB and other branding files at most 2 MB.|Colors must be #RRGGBB and primary color must meet 4.5:1 contrast against white.|Preview, remove/reset, and failed-upload behavior must preserve a valid configuration.")
web("System Settings - Login-page customization", [DB, SUPER_ACCOUNT],
    "Configured login heading, supporting text, logo, and background must appear in live preview and authentication pages.|Empty required login title must be rejected.|Cancel or upload failure must not publish incomplete customization.")
web("System Settings - Feature, survey, and community settings", [DB, SUPER_ACCOUNT],
    "Survey, Job Support, Community Forum, Messaging, and Notifications toggles must persist.|Survey title/instructions/availability and community guidelines/media-upload settings must save.|Dependent Messaging must not become usable when Community Forum is unavailable.|Disabled features must be reflected in affected user interfaces.")
web("System Settings - Maintenance mode", [DB, SUPER_ACCOUNT],
    "Enabling or disabling maintenance must require confirmation.|The selected state must persist.|Super Admin access must remain available while other protected roles are blocked.|Cancel must not change the maintenance state.")
web("Public Content Settings - About page", [DB, SUPER_ACCOUNT],
    "Existing About sections, title, description, CTA label, and image must load.|Valid JPG/PNG/WebP up to 4 MB must preview and save.|Invalid/oversized images and required empty text must be rejected.|Remove, restore, cancel, and save must affect only the intended draft/published state.")
web("Public Content Settings - FAQ editor", [DB, SUPER_ACCOUNT],
    "FAQ categories/items must support add, edit, delete, reorder, and visibility changes.|Category name, question, and answer must be non-empty.|Limits of 50 categories and 100 items per category must be enforced.|Preview/cancel/save must preserve order and visibility accurately.")
web("Public Content Settings - Privacy Policy editor", [DB, SUPER_ACCOUNT],
    "Introduction, dates, and sections must support add, edit, delete, reorder, visibility, preview, cancel, and save.|Required content and valid dates must be enforced.|The 100-section limit must be enforced.|Only allowed p, strong/b, ul/ol/li, h2/h3, and br markup must remain after sanitization.")

# Alumni Admin workflows
web("Alumni Verification - Account review queue", [DB, ALUMNI_ADMIN_ACCOUNT, "Pending, approved, and rejected Graduate accounts must exist."],
    "Pending/approved/rejected tabs, search, counts, refresh, and empty states must work.|The detail view must show survey identity, account status, registry evidence, and prior review data.|Only Alumni Admin must access the review actions.")
web("Alumni Verification - Approve or reject account", [DB, ALUMNI_ADMIN_ACCOUNT, "A pending Graduate account must exist."],
    "Approval must require confirmation and grant eligible portal access.|Reject must accept an optional reason and set rejected verification status.|A completed decision must update reviewer/date/status once.|Canceled or failed review must leave the prior status.")
web("Alumni Registry - Search, filters, sort, and pagination", [DB, ALUMNI_ADMIN_ACCOUNT, "Official alumni registry records must exist."],
    "Name search, course, batch, registry-status, and survey-answer filters must work.|Sortable columns and pagination must retain the current query.|No-result and reset behavior must be clear.|Summary/filter values must match stored registry records.")
web("Alumni Registry - View and edit record", [DB, ALUMNI_ADMIN_ACCOUNT, "An imported registry record must exist."],
    "Details must show name, course, batch, registry status, import source/date, and linked-account data.|Name, course, four-digit batch, and status edits must validate and save.|Cancel or failed save must preserve the previous record.")
web("Alumni Registry - Verify, inactivate, link, and unlink", [DB, ALUMNI_ADMIN_ACCOUNT, "Registry and Graduate account records must exist."],
    "Verify and Inactive actions must require confirmation and update the correct record.|Link search must show candidates and match strength, then link only the selected account.|Unlink must require confirmation and remove only the relationship.|Conflicting links or unauthorized calls must be rejected.")
web("Alumni Registry - Delete record", [DB, ALUMNI_ADMIN_ACCOUNT, "A deletable registry record must exist."],
    "The selected alumni name must appear in the confirmation.|Cancel must preserve the record.|Confirmed delete must remove only the registry record and safely handle linked-data restrictions or server errors.")
web("Alumni Registry - XLSX/CSV import preview", [DB, ALUMNI_ADMIN_ACCOUNT, "A registry XLSX or CSV file must be available."],
    "Only XLSX/CSV files up to 10 MB must be accepted.|Worksheet selection and preview must report detected, valid, duplicate, invalid, and ignored rows.|Unsupported/oversized/corrupt files must be rejected without writing data.|Duplicate handling must offer skip, update, or cancel.")
web("Alumni Registry - Import save and error report", [DB, ALUMNI_ADMIN_ACCOUNT, "A completed registry import preview must exist."],
    "Confirmation must identify importable valid rows.|The result must report processed, imported, updated, skipped-duplicate, invalid, and error totals.|Cancel duplicate behavior must disable saving.|Import issues must be downloadable as a CSV error report.")
web("Alumni Registry - CSV/XLSX export", [DB, ALUMNI_ADMIN_ACCOUNT, "Registry records must exist."],
    "Filtered, all, course, batch, unclaimed, and registered scopes must export correctly.|Course/batch scopes must require a selection.|CSV and XLSX formats must download with safe cell values.|An empty/error result must be reported without a misleading file.")
web("Announcement Management - List, filters, create, and edit", [DB, ALUMNI_ADMIN_ACCOUNT],
    "Draft, published, and archived records must filter correctly.|Title, summary, content, category, optional event date, and status must save with required and length validation.|Edit must prefill the selected record.|Cancel or API failure must preserve stored content.")
web("Announcement Management - Cover and gallery media", [DB, ALUMNI_ADMIN_ACCOUNT, "An editable announcement must exist."],
    "JPG, PNG, WebP, or GIF files up to 5 MB must be accepted.|Up to ten additional gallery images must be supported with preview and removal.|Invalid, oversized, or excess media must be rejected.|Published media must appear on the appropriate announcement interface.")
web("Announcement Management - Publish/archive/delete controls", [DB, ALUMNI_ADMIN_ACCOUNT, "Announcement records in multiple statuses must exist."],
    "Status changes and deletion must act on the selected record and require intended confirmation where implemented.|Only published Alumni Admin announcements must reach the Graduate Portal.|Canceled/failed delete must preserve content.")
web("Forum Moderation - Review queue and filters", [DB, ALUMNI_ADMIN_ACCOUNT, "Pending, approved, hidden, and reported posts must exist."],
    "Summary counts and pending/approved/hidden/all filters must load.|Search and category filters must find matching title, content, category, or author.|Reports, media, comments, and author/date information must be visible to the moderator.|No-result and server-error states must be clear.")
web("Forum Moderation - Post status and removal", [DB, ALUMNI_ADMIN_ACCOUNT, "A moderatable post must exist."],
    "Approve, Mark Pending, and Hide must update the selected post status.|Hidden or pending posts must not appear as approved Graduate feed content.|Delete Post must require confirmation and remove the post/comments only after approval.|Failed actions must retain the prior status.")
web("Forum Moderation - Comment removal and report context", [DB, ALUMNI_ADMIN_ACCOUNT, "A post with comments and pending reports must exist."],
    "The moderator must see reporter, target type, reason, and report date.|Delete Comment must identify the commenter and require confirmation.|Cancel must preserve the comment.|Confirmed removal must affect the selected comment only.")
web("Job Approval - Search and status filters", [DB, ALUMNI_ADMIN_ACCOUNT, "Pending, approved, and declined job posts must exist."],
    "Pending/approved/declined/all filters and summary counts must work.|Search must match graduate, program, company, skills, or job title.|Details must show job requirements, application information, poster, active state, and prior review information.|No-result state must be displayed.")
web("Job Approval - Approve or decline job", [DB, ALUMNI_ADMIN_ACCOUNT, "A pending job post must exist."],
    "Optional review notes must be retained with the decision.|Approval or decline must require the intended confirmation and update reviewer/date once.|Approved active jobs must become browseable and declined jobs must remain hidden.|Cancel or server failure must preserve pending status.")

# Cross-cutting reliability and authorization
web("Cross-Role Authorization - Backend API enforcement", [DB, "Authenticated accounts for all implemented roles must exist."],
    "Every administrator and Graduate API must validate the session and allowed role.|Changing route parameters or request bodies must not bypass program scope, record ownership, or feature permissions.|Unauthorized responses must not include protected record data.")
web("Cross-Cutting Validation - Missing records and invalid identifiers", [DB, "Test URLs and API requests with unknown or malformed IDs must be prepared."],
    "Unknown announcements, profiles, surveys, posts, rooms, jobs, and administrator records must return a controlled not-found/unavailable response.|Malformed IDs must not cause unhandled errors or disclose SQL details.|No unrelated record may be modified.")
web("Cross-Cutting Reliability - API and database failure presentation", ["A test environment must allow a controlled API/database failure.", SERVER],
    "The active interface must show a concise error and retry path where supported.|Forms must retain unsaved user input when safe.|No duplicate or partial operation may be presented as successful.|Technical stack traces, SQL, and secrets must not be displayed.")
web("Cross-Cutting Integrity - Audit and related-record consistency", [DB, "Representative create, update, delete, approval, and sign-in actions must be executable."],
    "Audited administrative actions must identify actor, module, action, and available before/after values.|Linked Graduate, account, survey, registry, forum, and job records must remain consistent after successful operations.|Failed or canceled operations must not create misleading success audit entries.")


mobile_cases: list[tuple[str, list[str], list[str]]] = []


def mobile(description: str, prereq: str | Sequence[str], expected: str | Sequence[str]) -> None:
    mobile_cases.append((description, bullets(prereq), bullets(expected)))


mobile("Responsive Web - Public home and navigation", [DB, "A mobile browser or responsive emulator at 320-480 px width must be available."],
       "The page must fit without unintended horizontal scrolling.|The mobile navigation must open, close, and reach each public route.|Text, images, and calls to action must remain readable and tappable.|Orientation change must not break the layout.")
mobile("Responsive Web - Public announcements and details", [DB, "Published public announcements must exist.", "A mobile browser must be available."],
       "Announcement cards, search, filters, pagination, and details must fit the viewport.|Gallery images must scale without clipping.|Preview controls must be reachable by touch.|Back navigation must return to the previous list state.")
mobile("Responsive Web - About, FAQ, and Privacy pages", [DB, "Configured public content and a mobile browser must be available."],
       "Configured sections must stack in a readable order.|FAQ expand/collapse controls must be touch accessible.|Long privacy content and lists must wrap inside the viewport.|Images must retain aspect ratio.")
mobile("Responsive Web - Admin sign in and forgot password", [DB, "A mobile browser and administrator test account must be available."],
       "Inputs, show/hide, submit, forgot-password, OTP, and reset controls must fit without overlap.|The on-screen keyboard must not hide the active control permanently.|Validation and status messages must wrap and remain readable.")
mobile("Responsive Web - Graduate sign in and forgot password", [DB, "A mobile browser and Graduate test account must be available."],
       "Sign-in and three-step reset forms must remain usable at narrow widths.|OTP controls and resend state must be touch accessible.|Successful actions must navigate to the correct responsive page.|Errors must not overflow the viewport.")
mobile("Responsive Web - Survey verification", [DB, "An active survey and matching Graduate record must exist.", "A mobile browser must be available."],
       "Verification method, identity fields, program selection, and submit control must stack correctly.|Student-number and email validation must remain visible near the field.|Loading, error, and success states must not shift controls off-screen.")
mobile("Responsive Web - Multi-section tracer survey", [DB, "A valid unused survey token and mobile browser must be available."],
       "Consent, progress, questions, choice controls, Other text, and Previous/Next actions must be usable by touch.|Long labels and option lists must wrap.|Draft answers must remain after orientation change or refresh.|No required question must be hidden behind fixed controls.")
mobile("Responsive Web - PSGC address controls", [DB, "PSGC data and a valid survey token must be available."],
       "Region, province, city/municipality, and barangay selectors must fit the viewport.|Dependent options must update without page-width overflow.|Mobile keyboard/select behavior must not lose already chosen address values.")
mobile("Responsive Web - Graduate announcements and dashboard", [DB, GRAD_ACCOUNT, "A mobile browser must be available."],
       "Portal navigation and feature tiles must stack and remain tappable.|Announcement search/filter/detail and dashboard indicators must fit without clipped text.|The mobile menu must show only enabled Graduate modules.")
mobile("Responsive Web - Graduate own/public profile", [DB, GRAD_ACCOUNT, "A mobile browser and another Graduate profile must be available."],
       "Cover, avatar, sections, and editing controls must adapt to narrow screens.|Upload controls and validation must be usable by touch.|Private controls must remain absent from another Graduate's profile.|Long email/address text must wrap.")
mobile("Responsive Web - Community Forum feed", [DB, GRAD_ACCOUNT, "Forum must be enabled with posts and a mobile browser."],
       "Post cards, filters, action buttons, comments, and counts must stack without overlap.|Like, comment, report, profile, and pagination actions must have usable touch targets.|Long content and media must remain inside the viewport.")
mobile("Responsive Web - Create/edit forum post and media viewer", [DB, GRAD_ACCOUNT, "Forum media uploads must be enabled on a mobile browser."],
       "Title/content/category and mobile file picker must remain usable.|Media previews and removal controls must fit.|Image/video viewer, next/previous, zoom where supported, and close must respond to touch.|Cancel must return without saving unintended changes.")
mobile("Responsive Web - Direct messaging", [DB, GRAD_ACCOUNT, "Messaging and Socket.IO must be enabled on a mobile browser."],
       "Conversation list, directory, active room, composer, send, and attachment actions must be navigable at narrow width.|The keyboard must not permanently cover the composer or latest message.|Unread, typing, delivery, presence, and connection states must remain visible.")
mobile("Responsive Web - Group-chat creation", [DB, GRAD_ACCOUNT, "At least two selectable Graduate accounts and a mobile browser must be available."],
       "Member search/filter, selection chips/list, group name, cancel, and create controls must fit the screen.|Scrolling the candidate list must not lose selections.|Validation must remain readable above the on-screen keyboard.")
mobile("Responsive Web - Message attachment and image preview", [DB, GRAD_ACCOUNT, "A mobile chat room and supported test files must be available."],
       "The device file picker must accept supported files and reject invalid/oversized ones.|Image messages must scale within the chat bubble.|Full preview must support open and close without trapping navigation.|Upload progress/failure must not cover the composer.")
mobile("Responsive Web - Browse jobs and job details", [DB, GRAD_ACCOUNT, "Approved jobs and Job Support must be enabled on a mobile browser."],
       "Search/filter controls and job cards must stack correctly.|Job detail content and external application links/files must be readable and tappable.|Long descriptions, skill lists, and salary/location text must wrap.|Closing detail must restore the list position.")
mobile("Responsive Web - Graduate job-posting form", [DB, GRAD_ACCOUNT, "The Graduate must have posting permission and use a mobile browser."],
       "Required and optional fields must fit and use appropriate mobile inputs.|Validation must remain associated with the field.|Create/edit/cancel/delete confirmation controls must be reachable.|Submission status must not be hidden by the keyboard.")
mobile("Responsive Web - Notifications", [DB, "An authenticated user with notifications must use a mobile browser."],
       "The notification panel must remain within the viewport and be scrollable.|Unread count, individual/all read actions, and destination links must be touch accessible.|The panel must close without blocking the underlying page.")
mobile("Responsive Web - Administrator sidebar and common controls", [DB, "An administrator must be authenticated on a mobile browser."],
       "The sidebar must open/close and show only authorized links.|Profile, theme, notifications, assistant where available, and logout must remain accessible.|Content must not remain obscured after navigation.")
mobile("Responsive Web - Admin dashboard, participation, and Dean view", [DB, "An Admin or Dean account and report data must exist.", "A mobile browser must be available."],
       "Cards and charts must scale or provide controlled horizontal scrolling only inside their containers.|Filters, selection, notifications, answer detail, and pagination must remain usable.|Dean program scope must remain unchanged on mobile.")
mobile("Responsive Web - Registrar graduate management", [DB, REGISTRAR_ACCOUNT, "A mobile browser and Graduate records must be available."],
       "Program tabs, filters, record cards/table fallback, pagination, and selection must fit.|Add/edit validation and confirmation dialogs must be usable with the keyboard.|Import file selection and result summary must not overflow.")
mobile("Responsive Web - Survey management and reports", [DB, ADMIN_ACCOUNT, "A mobile browser and survey/report data must be available."],
       "Survey cards, editor controls, question options, previews, response details, and analytics must remain operable.|Report tabs, filters, charts/tables, and export actions must fit or scroll inside defined containers.|No page-level horizontal overflow may hide actions.")
mobile("Responsive Web - Super Admin user management and settings", [DB, SUPER_ACCOUNT, "A mobile browser must be available."],
       "User filters, account forms, reminder tabs, audit details, backup controls, and settings tabs must be reachable.|Asset previews, color inputs, feature toggles, editors, dialogs, and save/cancel actions must adapt to narrow width.|Destructive confirmations must not extend off-screen.")
mobile("Responsive Web - Alumni Admin workflows", [DB, ALUMNI_ADMIN_ACCOUNT, "A mobile browser and representative alumni, announcement, forum, and job records must exist."],
       "Verification, registry, import/export, announcement editor, moderation, and job-approval views must offer a usable stacked/card or controlled-scroll layout.|Review notes, reasons, media, dialogs, filters, and actions must remain readable and tappable.")
mobile("Responsive Web - Rotation, resize, and session continuity", [DB, "An authenticated mobile-browser session must be active."],
       "Portrait-to-landscape and viewport resize must not log the user out or lose safe unsaved form state.|Menus and dialogs must recalculate inside the viewport.|Returning from background must restore the permitted page or require sign-in when the session expired.")
mobile("Responsive Web - Mobile accessibility and browser compatibility", [DB, "Current Chrome/Edge mobile emulation and at least one WebKit-based mobile browser must be available."],
       "Controls must be reachable by keyboard/screen-reader focus where applicable and have visible labels.|Text at 200% zoom must remain usable without cut-off.|Supported browsers must preserve core authentication, survey, portal, and administrator functions.|Unsupported device features must fail with clear alternatives.")


black_box_groups: list[tuple[str, list[tuple[str, str]]]] = []


def group(name: str, cases: Sequence[tuple[str, str]]) -> None:
    black_box_groups.append((name, list(cases)))


group("Public Home and Navigation Interface", [
    ("The visitor opens the GradTrack home page.", "The configured system branding, public content, and available calls to action are displayed."),
    ("The visitor uses each desktop or mobile navigation link.", "Each link opens its corresponding public page without an invalid route."),
    ("The visitor selects the survey call to action.", "The system opens Survey Verification, not a protected administrator or Graduate page."),
    ("The visitor opens an unknown public route.", "A controlled fallback/not-found experience is shown without application failure."),
    ("The public-content API is temporarily unavailable.", "The page shows a safe error or fallback and does not expose technical details."),
])
group("Public Announcements Interface", [
    ("The visitor opens the public announcement list with published Admin announcements available.", "Only eligible published public announcements are displayed."),
    ("The visitor searches for a matching announcement title or content.", "Matching announcements are returned and pagination starts from the first result page."),
    ("The visitor searches for text with no matching announcement.", "A clear no-results state is displayed."),
    ("The visitor filters by an announcement category.", "Only published announcements in that category are displayed."),
    ("The visitor opens a valid announcement.", "The correct announcement details and available media are displayed."),
    ("The visitor opens a missing, draft, or unpublished announcement ID.", "The announcement is unavailable and unpublished content is not disclosed."),
])
group("About, FAQ, and Privacy Interfaces", [
    ("The visitor opens About with configured visible sections.", "Visible sections appear in saved order with valid image and call-to-action content."),
    ("The visitor expands and collapses an FAQ item.", "The selected answer opens and closes without changing another item's content."),
    ("The visitor opens FAQ when no visible items exist.", "A clean empty state is displayed."),
    ("The visitor opens the Privacy Policy.", "The introduction, dates, and visible sanitized sections are displayed."),
    ("Stored public content contains disallowed script markup.", "Unsafe markup is not executed or rendered as active script."),
])
group("Admin Login Interface", [
    ("An active administrator enters valid email and password.", "The account is authenticated and redirected to the correct role landing page."),
    ("An administrator enters an incorrect password.", "A generic invalid-credentials message is displayed and no session is created."),
    ("The form is submitted with empty required fields.", "Required-field validation is displayed."),
    ("The user enters an invalid email format.", "Email-format validation prevents a successful submission."),
    ("An inactive administrator enters correct credentials.", "Access is denied because the account is inactive."),
    ("A non-Super-Admin signs in during maintenance mode.", "The user is blocked by the maintenance interface."),
    ("The user toggles password visibility.", "The password is shown or masked without changing its value."),
])
group("Graduate Login Interface", [
    ("An approved active Graduate enters valid credentials.", "The Graduate is authenticated and redirected to Graduate Announcements."),
    ("A Graduate enters an incorrect password.", "A generic invalid-credentials message is displayed."),
    ("The login form is submitted empty.", "Required email and password messages are displayed."),
    ("A pending-verification Graduate enters correct credentials.", "Access is denied and the user is informed that alumni verification is pending."),
    ("A rejected Graduate enters correct credentials.", "Access is denied with the implemented rejected-account guidance."),
    ("An inactive Graduate enters correct credentials.", "No Graduate Portal session is granted."),
    ("A Graduate attempts login during maintenance mode.", "The maintenance page is displayed."),
    ("The Graduate toggles password visibility.", "The control changes masking without modifying the password."),
])
group("Administrator Forgot Password Interface", [
    ("The user requests an OTP for a registered administrator email.", "A generic success response is shown and an eligible account receives a six-digit OTP."),
    ("The user requests an OTP for an unregistered email.", "The same generic response is shown without revealing account existence."),
    ("The user enters an invalid email format or leaves email empty.", "The request is rejected by form validation."),
    ("The user enters a valid unexpired OTP.", "The reset-password step is opened."),
    ("The user enters an expired, malformed, or incorrect OTP.", "Verification is denied with an appropriate error."),
    ("The user exceeds five invalid OTP attempts.", "Further use of that OTP flow is blocked as implemented."),
    ("The user requests resend before one minute.", "Resend remains unavailable during the cooldown."),
    ("The user enters a weak or mismatched new password.", "Password complexity or matching validation is displayed."),
    ("The user submits a valid complex matching password.", "The password is reset, the reset token is invalidated, and Admin Sign In is shown."),
])
group("Graduate Forgot Password Interface", [
    ("The Graduate requests an OTP for a registered email.", "A generic success response is shown and an eligible Graduate receives a six-digit OTP."),
    ("The Graduate requests an OTP for an unknown email.", "The response does not disclose whether the account exists."),
    ("The Graduate submits empty or invalid email.", "Email validation is displayed."),
    ("The Graduate submits a valid unexpired OTP.", "The password-reset step opens."),
    ("The Graduate submits an expired or incorrect OTP.", "OTP verification is denied."),
    ("The Graduate exceeds the invalid-attempt limit.", "The current OTP flow can no longer be used."),
    ("The Graduate requests OTP resend during cooldown.", "The resend action remains disabled until allowed."),
    ("The Graduate submits weak or mismatched passwords.", "The new password is not saved and validation is displayed."),
    ("The Graduate submits a valid complex matching password.", "The password changes and the Graduate returns to Sign In."),
])
group("Role Routing and Session Interface", [
    ("Each implemented administrator role signs in.", "Admin, Registrar, Alumni Admin, Super Admin, Dean, MIS Staff, and Research Coordinator reach their implemented landing pages."),
    ("An unauthenticated visitor opens a protected administrator URL.", "The visitor is redirected to Admin Sign In."),
    ("An unauthenticated visitor opens a protected Graduate URL.", "The visitor is redirected to Graduate Sign In."),
    ("A signed-in role opens another role's route directly.", "The route is denied and a safe role-specific fallback is used."),
    ("A user refreshes a protected page with a valid session.", "The permitted page remains available after session validation."),
    ("A user confirms logout and then uses browser Back.", "Protected content is not restored without authentication."),
    ("A user cancels logout.", "The current session and page remain active."),
])
group("Survey Verification Interface", [
    ("A matching Graduate verifies with a correctly formatted student number, last name, and program.", "A survey token is issued and the active questionnaire opens."),
    ("A matching Graduate verifies with registered email, last name, and program.", "A survey token is issued for the correct Graduate and survey."),
    ("The user enters a student number outside the 2XXX-XXXX format.", "Format validation is displayed and verification does not proceed."),
    ("The user enters a wrong last name or program.", "The identity is not verified."),
    ("The user leaves a required identity field empty.", "Required-field messages are displayed."),
    ("No Registrar record matches the entered identity.", "A clear no-match response is shown without exposing another record."),
    ("No survey is active.", "The interface reports that no active survey is available."),
    ("The Graduate already submitted the active survey.", "A second survey token/submission is blocked."),
    ("The user opens an expired or malformed survey token.", "The user is returned to verification with an invalid/expired message."),
    ("The verification API fails.", "A recoverable error is shown and no token is falsely issued."),
])
group("Tracer Survey Interface", [
    ("The verified user attempts to start without accepting the privacy agreement.", "The questionnaire remains locked until consent is accepted."),
    ("The user answers valid values for every required question.", "Each section validates and the user may reach final submission."),
    ("The user leaves a required answer empty and selects Next.", "The section does not advance and the missing answer is identified."),
    ("The user selects Other without supplying supporting text.", "The Other response is rejected until text is entered."),
    ("The user changes an employment answer that controls conditional questions.", "Only relevant branch questions remain active and hidden answers do not block progress."),
    ("The user selects region, province, city/municipality, and barangay in order.", "Each dependent PSGC list contains values for its parent selection."),
    ("The user submits an incomplete or inconsistent required address.", "Address validation prevents submission."),
    ("The user uses Previous and Next after entering answers.", "Entered answers remain associated with the correct questions."),
    ("The user refreshes a valid partially completed survey.", "The saved draft can be restored for the same token."),
    ("The user submits a fully valid survey.", "One response is stored, the token becomes submitted, and success is shown."),
    ("The user submits the same token again.", "No duplicate response is created."),
    ("The server fails during submission.", "The user sees an error and retained draft data; no false success is shown."),
    ("The active survey contains each supported question type.", "Text, date, choice, checkbox, radio, rating, and header items render and save correctly."),
])
group("Graduate Account Registration Interface", [
    ("A survey completer enters a unique valid email and complex matching password.", "One pending-verification Graduate account is created and approval guidance is shown."),
    ("The user enters an invalid email.", "Email-format validation prevents account creation."),
    ("The user enters a weak password.", "The complexity requirements are displayed and no account is created."),
    ("The password confirmation differs.", "Mismatch validation is displayed."),
    ("The email is already registered.", "Duplicate account creation is rejected."),
    ("The Graduate is already linked to an account.", "A second account is not created."),
    ("The response ID does not belong to the Graduate identity.", "The request is rejected by the server."),
    ("The registration API fails.", "No partially linked account is reported as successful."),
])
group("Graduate Announcements Interface", [
    ("The Graduate opens Announcements after sign-in.", "Published Alumni Admin announcements are displayed."),
    ("The Graduate searches and filters by category.", "Only matching published announcements are returned."),
    ("The Graduate searches with no result.", "A clear empty state and filter-reset option are shown."),
    ("The Graduate opens a valid announcement with multiple images.", "Details and the correct gallery are displayed."),
    ("The Graduate uses next, previous, and close in image preview.", "The gallery changes images and closes correctly."),
    ("The Graduate opens a missing or unpublished record.", "The announcement is unavailable and hidden content is not disclosed."),
])
group("Graduate Dashboard Interface", [
    ("A Graduate with complete survey/profile data opens the dashboard.", "The correct alumni score, employment/alignment indicators, badges, and tiles are displayed."),
    ("A Graduate with missing optional data opens the dashboard.", "The dashboard uses safe empty/default states rather than failing."),
    ("The Graduate selects an enabled module tile.", "The matching portal module opens."),
    ("The Graduate selects or directly requests a disabled feature.", "The defined unavailable state is shown and no hidden feature action is executed."),
])
group("Graduate Profile Interface", [
    ("The Graduate opens their own profile.", "Editable basic details and read-only survey-derived sections are displayed."),
    ("The Graduate saves valid name, email, phone, and address changes.", "The profile is updated and shows the stored values after refresh."),
    ("The Graduate submits an invalid or duplicate email.", "The update is rejected and the previous email remains."),
    ("The Graduate opens another alumni profile.", "Only the implemented public information is shown and edit/security controls are absent."),
    ("The Graduate uploads a valid JPG, PNG, WebP, or GIF profile image up to 5 MB.", "The new avatar is stored and displayed."),
    ("The Graduate uploads an invalid type or image over 5 MB.", "The upload is rejected and the current avatar remains."),
    ("The Graduate removes the profile image.", "The default avatar is restored."),
    ("The Graduate uploads a valid cover photo.", "The cover is stored and displayed after refresh."),
    ("The Graduate uploads an invalid or oversized cover.", "The upload is rejected without replacing the prior cover."),
    ("The Graduate removes the cover photo.", "The default cover presentation returns."),
    ("The Graduate changes password using the correct current password and valid new password.", "The new password is saved and the old password no longer authenticates."),
    ("The Graduate enters a wrong current password, weak password, or mismatch.", "The password is not changed and the appropriate validation is displayed."),
    ("The Graduate opens an unknown profile ID.", "A controlled unavailable profile state is shown."),
])
group("Community Forum Interface", [
    ("The Graduate opens the forum with approved posts available.", "Approved posts, authors, categories, counts, and media are displayed."),
    ("The Graduate searches and applies category, program, and year filters.", "Only posts matching all selected criteria are displayed."),
    ("The Graduate uses filters that match no post.", "A clear no-results state is displayed."),
    ("The Graduate creates a valid text post.", "The post is saved for the authenticated Graduate with the implemented moderation status."),
    ("The Graduate submits a post with empty required fields.", "Required-field messages prevent creation."),
    ("The Graduate submits prohibited content.", "The moderation validation rejects or flags the content as implemented."),
    ("The Graduate uploads up to ten valid image/video files.", "Valid media is previewed, uploaded, and associated with the post."),
    ("The Graduate uploads excess, invalid, or oversized media.", "The media is rejected with the applicable limit message."),
    ("The owner edits an existing post and saves valid changes.", "The post is updated and returned to pending moderation as implemented."),
    ("The owner cancels editing.", "No post data changes."),
    ("The owner confirms deletion.", "The post and its visible related content are removed."),
    ("A Graduate attempts to edit or delete another Graduate's post.", "The action is not available and the API denies it."),
    ("The Graduate likes and then unlikes a post.", "The Graduate's reaction and count update once for each state."),
    ("The Graduate submits a valid comment and then a blank comment.", "The valid comment is saved; the blank comment is rejected."),
    ("The Graduate reports a post or comment with a reason.", "A moderation report is created for the selected target."),
    ("The Graduate opens and closes forum media preview.", "The correct media opens and the viewer returns to the same post when closed."),
])
group("Messaging Interface", [
    ("The Graduate searches the alumni directory by name, program, or batch.", "Matching eligible Graduate accounts are displayed."),
    ("The Graduate selects one alumnus for direct messaging.", "The existing direct room opens or one direct room is created."),
    ("The Graduate creates a group with a name and at least two members.", "A group room with the selected members is created."),
    ("The Graduate submits a group with empty name or too few members.", "Group creation is rejected with validation."),
    ("The Graduate sends a non-empty text message under 5,000 characters.", "The message progresses from pending to sent and appears in the active room."),
    ("The Graduate sends blank text or text over 5,000 characters.", "The message is not sent."),
    ("Another room member sends a message while both users are connected.", "The message appears in real time without refresh."),
    ("The Graduate uploads a supported image up to 10 MB.", "The image appears as its own message item for room members."),
    ("The Graduate opens and closes an image message.", "The full image preview opens and closes without altering the message."),
    ("The Graduate uploads a supported PDF/Office/text/CSV file up to 25 MB.", "A usable attachment message with the original filename is sent."),
    ("The Graduate uploads an oversized or dangerous executable/script file.", "The attachment is rejected."),
    ("The Graduate uploads a file with a mismatched MIME type.", "The server rejects the upload."),
    ("A participant types and then stops.", "Other participants see a temporary typing indicator that clears appropriately."),
    ("A participant connects and disconnects.", "Online/offline state and last-seen information update correctly."),
    ("The recipient opens the conversation containing a delivered message.", "The correct delivery/read receipt and unread count are updated."),
    ("The Graduate loads older messages.", "Older items are prepended in order without duplication."),
    ("A new message arrives while the Graduate is viewing older history.", "A new-message indicator is shown without moving the reading position unexpectedly."),
    ("The realtime connection drops while a message is sent.", "The message is marked failed/pending and retry is offered rather than false success."),
    ("The realtime connection reconnects after messages were missed.", "Missed messages synchronize once, in order, with correct unread state."),
    ("A Graduate attempts to access a room they do not belong to.", "The room and its messages are denied."),
])
group("Browse Jobs Interface", [
    ("The Graduate opens Browse Jobs with approved active posts available.", "Only approved active jobs are listed."),
    ("The Graduate searches by title, company, skills, location, or program fit.", "Matching jobs are displayed."),
    ("The Graduate searches with no match.", "A clear no-results state is displayed."),
    ("The Graduate opens a valid job detail.", "The correct job, poster, dates, requirements, and available application information are displayed."),
    ("A job has a contact email, external link, method, or requirements file.", "The configured external application information is usable."),
    ("The Graduate tries to open a missing job.", "A controlled unavailable response is shown."),
    ("A pending, declined, inactive, or expired post is requested.", "The post is not presented as an active browseable opportunity."),
    ("Job Support is disabled by Super Admin.", "Browse Jobs is hidden or shows the defined unavailable state."),
])
group("Graduate Job Posting Interface", [
    ("An eligible employed Graduate opens Job Posting.", "The create form and own-post list are available."),
    ("An ineligible Graduate opens Job Posting directly.", "The posting function remains locked and the API denies creation."),
    ("The Graduate submits title, company, description, and valid application contact/method.", "A pending-approval job post is created."),
    ("The Graduate omits a required title, company, or description.", "Required-field validation prevents creation."),
    ("The Graduate provides no application email, link, or method.", "The post is rejected because application information is required."),
    ("The Graduate enters an invalid application email.", "Email validation prevents submission."),
    ("The Graduate saves valid optional job type, location, industry, salary, skills, program fit, and deadline.", "The optional values appear in the owner's detail and later approved detail."),
    ("The owner edits an approved or pending job.", "Valid changes save and the post returns to pending approval."),
    ("The owner cancels editing.", "The stored post remains unchanged."),
    ("The owner confirms job deletion.", "The selected own job is removed."),
    ("The Graduate attempts to edit/delete another Graduate's job.", "The action is denied."),
    ("The approval service fails during create/edit.", "No duplicate or false successful post is shown."),
])
group("Notifications Interface", [
    ("An authenticated user has unread notifications.", "The bell shows the correct unread count and relevant notification items."),
    ("The user marks one notification read.", "Only that item's read state and the unread count change."),
    ("The user marks all notifications read.", "All visible notifications become read and the count becomes zero."),
    ("The user selects a notification with a permitted destination.", "The corresponding record/interface opens."),
    ("The user has no notifications.", "A clear empty state is displayed."),
    ("Notifications are disabled or the API fails.", "The bell is hidden/unavailable as configured or a safe error is shown without blocking the current page."),
])
group("Administrator Profile and Layout Interface", [
    ("An administrator opens Profile.", "The correct name, read-only email, role, and current image are displayed."),
    ("The administrator saves a valid full-name change.", "The updated name is stored and reflected in the layout."),
    ("The administrator uploads a valid image up to 5 MB.", "The profile image is stored and displayed."),
    ("The administrator uploads an invalid or oversized image.", "The upload is rejected and the prior image remains."),
    ("The administrator changes password using valid current/new values.", "The credential is updated and the old password stops working."),
    ("The administrator enters wrong current password, weak new password, or mismatch.", "The password is not changed."),
    ("The administrator toggles the theme.", "The interface theme changes consistently without changing authorization."),
    ("The administrator cancels logout.", "The session remains active."),
    ("The administrator confirms logout.", "The session ends and protected pages require sign-in."),
])
group("Admin Dashboard Interface", [
    ("The Admin opens the dashboard with reportable data.", "Survey snapshot, totals, employability, trend, and job-alignment widgets load with consistent values."),
    ("The Admin selects another survey, year, or program.", "All affected dashboard values update to the selected scope."),
    ("The selected scope contains no data.", "Zero/empty states are shown without invented percentages."),
    ("The dashboard API fails.", "An error/retry state is shown instead of stale success data."),
    ("MIS Staff or Research Coordinator opens the implemented dashboard.", "The dashboard loads while unauthorized survey/report/CRUD navigation remains absent."),
])
group("Graduate Participation Interface", [
    ("The Admin selects a survey and filters Answered records.", "Only graduates with submitted responses for that survey are shown."),
    ("The Admin filters Not Answered records.", "Only eligible non-respondents are shown."),
    ("The Admin searches by name, email, or student ID with program/year filters.", "Matching participation records are displayed."),
    ("The filters produce no records.", "A no-results state is shown."),
    ("The Admin selects visible rows.", "Only visible eligible rows become selected."),
    ("The Admin selects all matching rows across pages.", "The selection represents the complete filtered eligible set."),
    ("The Admin attempts Notify with no eligible selection.", "Sending is prevented with guidance."),
    ("The Admin confirms Notify for eligible non-respondents.", "The result reports sent, failed, and skipped recipients."),
    ("The Admin opens answers for a completed respondent and then a non-respondent.", "Correct stored answers open for the respondent; no fabricated answers open for the non-respondent."),
])
group("Registrar Graduate Records Interface", [
    ("The Registrar selects a program tab and graduation year.", "Only matching Graduate records are displayed."),
    ("The Registrar searches by name, student ID, or email.", "Matching records are returned with filters preserved."),
    ("The Registrar searches with no match.", "A clear empty state is displayed."),
    ("The Registrar adds a record with valid required data.", "One Graduate record is saved under the correct program and year."),
    ("The Registrar leaves first/last name, program, year, or phone empty.", "Required-field validation prevents saving."),
    ("The Registrar enters a student ID outside dddd-dddd.", "Student-ID format validation is displayed."),
    ("The Registrar enters a phone not consisting of 11 digits beginning 09.", "Phone validation prevents saving."),
    ("The Registrar enters an invalid optional email.", "Email-format validation prevents saving."),
    ("The Registrar enters an existing student ID or email.", "The duplicate record is rejected."),
    ("The Registrar edits an existing record with valid values.", "The selected record is updated and relationships remain consistent."),
    ("The Registrar cancels an edit.", "The record remains unchanged."),
    ("The Registrar confirms single-record deletion.", "Only the selected record is deleted."),
    ("The Registrar cancels single or bulk deletion.", "No records are deleted."),
    ("The Registrar confirms selected-record deletion.", "Only selected records are deleted and the count updates."),
    ("The Registrar confirms deletion for one program/year cohort.", "Only that cohort is removed."),
    ("The Registrar imports a valid XLS/XLSX file.", "Valid rows are mapped and imported with a result summary."),
    ("The import contains missing required values, duplicates, or invalid rows.", "Invalid/skipped rows are identified and valid handling follows implementation without false counts."),
])
group("Dean Survey Participation Interface", [
    ("Dean-CCS opens Survey Participation.", "Only BSCS and ACT participation data is available."),
    ("Dean-COED opens Survey Participation.", "Only BSED and BEED participation data is available."),
    ("Dean-HM opens Survey Participation.", "Only BSHM participation data is available."),
    ("A Dean applies search, year, and response-status filters.", "Only matching records inside the Dean's program scope are displayed."),
    ("A Dean attempts to request another program through a modified URL/API parameter.", "The server denies or excludes out-of-scope records."),
    ("A Dean selects eligible non-respondents and confirms Notify.", "Email results are returned only for allowed programs."),
    ("A Dean opens a completed response in scope.", "The correct response detail is displayed."),
    ("A Dean tries to open an out-of-scope response.", "The response data is denied."),
])
group("Survey Management Interface", [
    ("The Admin opens Survey Management.", "Existing surveys and their draft/inactive/active status are displayed."),
    ("The Admin creates a survey with valid metadata when no active conflict exists.", "The tracer-template survey is created once."),
    ("The Admin leaves required survey metadata empty.", "Validation prevents survey creation."),
    ("The Admin attempts to create/activate a second active survey.", "The single-active-survey rule blocks the action."),
    ("The Admin edits title, description, or allowed status.", "The selected survey is updated."),
    ("The Admin adds a valid question and section.", "The items appear in the configured order."),
    ("The Admin configures text, date, multiple-choice, radio, rating, checkbox, and header items.", "Each type retains its appropriate configuration."),
    ("The Admin creates a choice question without required option data.", "Question validation prevents an unusable configuration."),
    ("The Admin changes required flags and question options.", "Preview and saved survey reflect the changes."),
    ("The Admin removes a question and cancels before save.", "The stored survey remains unchanged."),
    ("The Admin previews a configured survey.", "Sections, questions, options, and required markers match the editor."),
    ("The Admin confirms deletion of one survey.", "Only that survey and its defined dependent data are removed."),
    ("The Admin cancels survey deletion.", "The survey and responses remain."),
    ("The Admin invokes Clear All Surveys and cancels confirmation.", "No survey is removed."),
    ("The Admin confirms Clear All Surveys using the required confirmation flow.", "All surveys and defined dependent data are cleared, with a result/error message."),
    ("A non-Admin role opens Survey Management or calls its API.", "Access is denied."),
])
group("Survey Responses and Analytics Interface", [
    ("The Admin opens responses for a survey with submissions.", "The correct respondents and response records are listed."),
    ("The Admin opens one response detail.", "Identity and stored answers correspond to the selected response."),
    ("The Admin searches/filters to no response.", "A clean no-results state is displayed."),
    ("The Admin exports responses to CSV.", "A CSV for the selected survey/result set downloads with safe values."),
    ("The Admin opens analytics for a survey with data.", "Response rate, program/year/employment, and question analyses match stored records."),
    ("The Admin changes the selected survey.", "Every chart and summary refreshes to the new survey."),
    ("A survey has no responses or unsupported empty answers.", "Charts use a controlled zero/empty state."),
    ("The Admin exports survey analytics to Excel.", "The selected survey analytics are downloaded in an editable spreadsheet."),
    ("The Admin exports survey analytics to PDF.", "The selected survey analytics are downloaded in a readable PDF."),
])
group("Reports and Analytics Interface", [
    ("The Admin opens the Overview report.", "Summary, tables, and charts load from the same report scope."),
    ("The Admin combines employment, alignment, year, program, and location filters.", "Only records satisfying all selected filters contribute to results."),
    ("The Admin changes region, province, city/municipality, and barangay.", "Dependent choices and report scope update consistently."),
    ("The Admin resets filters.", "The unfiltered report is restored."),
    ("Filters return no records.", "Zero/empty outputs are displayed without invented values."),
    ("The Admin opens By Program.", "Program-based data and controls are displayed."),
    ("The Admin opens By Year.", "Graduation-year data and controls are displayed."),
    ("The Admin opens Employment Status and Salary Distribution.", "Employment and salary summaries use the selected scope."),
    ("The Admin opens Location.", "Geographic data is displayed from normalized stored location fields."),
    ("The Admin opens Survey Analytics from Reports.", "Survey-based metrics and questions correspond to the selected survey."),
    ("The Admin exports the filtered report to Excel.", "An Excel file with scope/filter metadata and intended tables downloads."),
    ("The Admin exports the filtered report to PDF.", "A readable PDF with intended summaries/charts and filter metadata downloads."),
    ("The Admin requests AI descriptive analysis and the service succeeds.", "Analysis based on the selected data is displayed."),
    ("The AI service times out or is not configured.", "A clear unavailable/retry state is shown and standard reports remain usable."),
])
group("User Management Interface", [
    ("The Super Admin searches and filters administrator accounts.", "Matching accounts are displayed and Graduate accounts remain excluded."),
    ("The Super Admin creates an account with valid username, email, password, name, role, and status.", "One administrator account is created."),
    ("A required account field is empty.", "Validation prevents creation."),
    ("The username already exists.", "Duplicate username is rejected."),
    ("The email already exists or has invalid format.", "The account is not created and the email error is displayed."),
    ("The password is invalid or weak.", "Password validation prevents creation/update."),
    ("An unsupported role value is submitted.", "The request is rejected."),
    ("The Super Admin edits valid account information.", "The selected account is updated."),
    ("The Super Admin leaves the optional new-password field empty during edit.", "Other changes save without changing the password."),
    ("The Super Admin cancels editing.", "No user data changes."),
    ("The Super Admin confirms deactivation of another administrator.", "The account becomes inactive and cannot sign in."),
    ("The Super Admin confirms activation of an inactive administrator.", "The account becomes active and may authenticate."),
    ("The logged-in Super Admin attempts to deactivate their own account.", "The action is blocked."),
    ("A non-Super-Admin calls User Management.", "Access is denied."),
    ("The user looks for a delete-user control.", "No delete control is exposed by the current User Management interface."),
])
group("Auto Email Reminders Interface", [
    ("The Super Admin opens reminders with an active survey.", "Active survey, eligible count, interval, statistics, and eligible-list control are displayed."),
    ("No survey is active.", "The interface instructs the user to activate a survey and Send is disabled."),
    ("No eligible non-respondent exists.", "Send is disabled and zero eligible recipients are shown."),
    ("Email delivery is disabled.", "The send action is prevented with the implemented status message."),
    ("The Super Admin confirms a manual reminder send.", "Eligible, sent, failed, and skipped totals are reported and attempts are logged."),
    ("The Super Admin cancels reminder confirmation.", "No reminder email is sent."),
    ("The Super Admin saves a preset interval.", "The selected interval is stored."),
    ("The Super Admin enters a custom value below 1 or above 365.", "The value is constrained/rejected to the supported range."),
    ("The Super Admin views reminder history.", "Auto/manual type, recipient, status, and date are displayed for logged attempts."),
    ("The scheduled job runs without the required secret.", "Automatic processing is denied and no unauthorized reminders are sent."),
])
group("Audit Trail Interface", [
    ("The Super Admin opens Audit Trail with logs available.", "Audit entries and pagination are displayed."),
    ("The Super Admin searches and combines role, department, action, module, and date filters.", "Only matching audit records are displayed."),
    ("The Super Admin clears filters.", "The default audit list is restored."),
    ("The Super Admin opens an audit detail.", "Available previous values, new values, and metadata are shown for that action."),
    ("The selected filter has no logs.", "A clean no-record state is displayed."),
    ("The Super Admin refreshes the list.", "Current logs are requested without losing intended filter state."),
    ("The Super Admin exports filtered logs to CSV.", "The CSV contains the intended filtered audit fields."),
    ("An exported value begins with =, +, -, or @.", "The value is neutralized against spreadsheet formula injection."),
    ("A non-Super-Admin requests audit data.", "Access is denied without returning logs."),
])
group("Database Backup Interface", [
    ("The Super Admin opens Database Backup.", "Database summary and table details are displayed."),
    ("The Super Admin refreshes backup summary.", "Current table, record, and size values are requested."),
    ("The Super Admin starts a backup while the database is available.", "A valid SQL download containing supported structure and data is produced."),
    ("The backup process encounters a database/server error.", "A clear failure is shown and no corrupt file is reported as successful."),
    ("A non-Super-Admin opens the backup route.", "The route is denied."),
    ("A non-Super-Admin calls the backup API directly.", "The API denies the request and does not expose database contents."),
    ("The database contains zero records in a table.", "The table remains represented safely without breaking the backup."),
])
group("System Settings Interface", [
    ("The Super Admin saves valid general system and institution values.", "The settings persist and appear in affected public/authentication interfaces."),
    ("Required system name, short name, institution, or login title is empty.", "Validation prevents saving."),
    ("The contact email has invalid format.", "Email validation prevents saving."),
    ("The Super Admin uploads a valid logo or login asset within its size limit.", "The asset previews, saves, and appears in the configured interface."),
    ("A non-favicon ICO or unsupported file is uploaded.", "The asset is rejected."),
    ("A branding file exceeds 2 MB or a background exceeds 5 MB.", "The oversized file is rejected."),
    ("The Super Admin enters a color outside #RRGGBB.", "Color validation prevents saving."),
    ("The primary color has less than 4.5:1 contrast against white.", "The color is rejected with contrast guidance."),
    ("The Super Admin removes or resets a branding asset and confirms save.", "The default/empty asset state is applied as intended."),
    ("The Super Admin cancels branding changes.", "The published branding remains unchanged."),
    ("The Super Admin changes login-page text and preview.", "The preview and saved login page reflect valid content."),
    ("The Super Admin disables Community Forum.", "Forum and dependent Messaging become unavailable to Graduates."),
    ("The Super Admin disables only Notifications or Job Support.", "Only the selected feature's Graduate controls become unavailable."),
    ("The Super Admin saves survey title/instructions/availability.", "Survey-facing configuration reflects the saved values."),
    ("The Super Admin saves community guidelines/default announcement/media setting.", "Community interfaces use the stored configuration."),
    ("The Super Admin enables maintenance and confirms.", "Non-Super-Admin protected access is blocked while Super Admin remains available."),
    ("The Super Admin cancels a maintenance change.", "The current maintenance state remains."),
    ("A non-Super-Admin calls the settings API.", "Access is denied."),
])
group("Public Website Content Settings Interface", [
    ("The Super Admin edits valid About title, description, sections, CTA label, and image.", "The saved visible About content appears on the public page in order."),
    ("A required About field is empty.", "The content is not saved and validation is displayed."),
    ("The About image is not JPG/PNG/WebP or exceeds 4 MB.", "The upload is rejected."),
    ("The Super Admin removes/restores an About image and previews changes.", "Preview reflects the draft and only Save publishes it."),
    ("The Super Admin cancels About edits.", "Published About content remains unchanged."),
    ("The Super Admin adds, edits, reorders, hides, and deletes valid FAQ content.", "Saved public FAQ order and visibility match the editor."),
    ("A FAQ category name, question, or answer is empty.", "Validation prevents saving the invalid item."),
    ("The editor reaches more than 50 FAQ categories or 100 items in a category.", "The configured limit prevents adding more."),
    ("The Super Admin cancels FAQ changes.", "Published FAQ content remains unchanged."),
    ("The Super Admin edits valid Privacy introduction, dates, and sections.", "The public Privacy Policy reflects saved visible content in order."),
    ("A Privacy required value is empty or a date is invalid.", "Validation prevents saving."),
    ("The editor reaches more than 100 Privacy sections.", "The configured section limit prevents adding more."),
    ("Privacy content includes allowed paragraph, bold, list, heading, and line-break markup.", "Allowed markup remains and renders safely."),
    ("Privacy content includes script or unsupported markup.", "Disallowed markup is removed or neutralized."),
    ("The Super Admin hides or reorders public sections and saves.", "The public page reflects visibility/order exactly."),
    ("The Super Admin cancels Privacy changes.", "Published Privacy content remains unchanged."),
])
group("Alumni Verification Interface", [
    ("The Alumni Admin opens the Pending review tab.", "Pending Graduate accounts and summary count are displayed."),
    ("The Alumni Admin searches by account identity.", "Matching review accounts are displayed."),
    ("The search or status tab has no accounts.", "A clear empty state is displayed."),
    ("The Alumni Admin opens account details.", "Identity, program/year, contact, survey, review, and registry-match information is shown."),
    ("The Alumni Admin confirms approval of a pending account.", "Verification becomes approved and eligible Graduate Portal access is activated."),
    ("The Alumni Admin cancels approval.", "The account remains pending."),
    ("The Alumni Admin rejects a pending account with an optional reason.", "Verification becomes rejected and the reason is stored/displayed."),
    ("The review API fails during a decision.", "The prior status remains and an error is displayed."),
    ("A non-Alumni-Admin attempts account review.", "The route/API denies access."),
])
group("Alumni Registry Interface", [
    ("The Alumni Admin searches and combines course, batch, registry-status, and survey-answer filters.", "Only matching official registry records are displayed."),
    ("The Alumni Admin changes sorting or page.", "The requested order/page loads while retaining filters."),
    ("Filters match no registry record.", "A clear no-results state is displayed."),
    ("The Alumni Admin opens a registry detail.", "Name, course, batch, status, import source/date, and linked-account information are displayed."),
    ("The Alumni Admin edits valid name, course, four-digit batch, and status.", "The selected registry record is updated."),
    ("The Alumni Admin enters empty/invalid edit values.", "Validation prevents the update."),
    ("The Alumni Admin cancels editing.", "The record remains unchanged."),
    ("The Alumni Admin verifies or marks a record inactive and confirms.", "The selected registry status changes."),
    ("The Alumni Admin searches account-link candidates.", "Candidates and match strength are displayed."),
    ("The Alumni Admin links one valid account candidate.", "The selected registry record and Graduate account become linked once."),
    ("The Alumni Admin attempts a conflicting/duplicate link.", "The server rejects the conflicting relationship."),
    ("The Alumni Admin confirms unlink.", "Only the relationship is removed."),
    ("The Alumni Admin confirms registry deletion.", "Only the selected registry record is removed."),
    ("The Alumni Admin cancels registry deletion.", "The record remains."),
    ("The Alumni Admin selects a valid XLSX or CSV up to 10 MB.", "Worksheets/rows are parsed and preview counts are displayed."),
    ("The file type is unsupported, corrupt, or over 10 MB.", "The file is rejected and no records are written."),
    ("Import preview contains duplicate and invalid rows.", "The preview separates valid, duplicate, invalid, and ignored counts/issues."),
    ("The Alumni Admin chooses Skip or Update duplicate handling and confirms.", "The result follows the selected behavior and reports imported/updated/skipped/invalid totals."),
    ("The Alumni Admin chooses Cancel duplicate handling.", "Save Import is disabled and nothing is written."),
    ("The Alumni Admin exports filtered/all/course/batch/unclaimed/registered scope as CSV or XLSX.", "A safe file containing the selected scope is downloaded; missing required scope selection is rejected."),
])
group("Announcement Management Interface", [
    ("The Alumni Admin filters draft, published, and archived announcements.", "Only records with the selected status are displayed."),
    ("The Alumni Admin creates an announcement with valid required content.", "The announcement is saved once with the selected category/status."),
    ("Title, summary, or content is empty.", "Required-field validation prevents saving."),
    ("Title exceeds 255, summary exceeds 500, or content exceeds 60,000 characters.", "The configured length limit prevents saving."),
    ("The Alumni Admin enters an optional valid event date.", "The date is stored and displayed with the announcement."),
    ("The Alumni Admin edits an existing announcement.", "The form is prefilled and valid changes save to the selected record."),
    ("The Alumni Admin cancels create/edit.", "No unintended record or change is saved."),
    ("A valid JPG/PNG/WebP/GIF cover up to 5 MB is uploaded.", "The cover previews and saves."),
    ("An invalid or oversized cover/gallery image is uploaded.", "The media is rejected."),
    ("Up to ten valid gallery images are selected.", "The images preview and save in the announcement gallery."),
    ("More than ten gallery images are selected.", "The excess selection is rejected."),
    ("The Alumni Admin removes a selected/saved media item.", "Only that media item is removed after save."),
    ("The announcement is published.", "It becomes visible in the Graduate Portal announcement interface."),
    ("The announcement is draft or archived.", "It is not visible in the Graduate Portal."),
    ("The Alumni Admin confirms deletion.", "The selected announcement and defined media records are removed."),
    ("The Alumni Admin cancels deletion or the API fails.", "The announcement remains stored."),
])
group("Forum Moderation Interface", [
    ("The Alumni Admin opens pending moderation.", "Pending posts, counts, author, content, media, comments, and reports are displayed."),
    ("The Alumni Admin searches by title/content/category/author and filters category/status.", "Only matching posts are displayed."),
    ("Filters produce no post.", "A clear no-results state is shown."),
    ("The Alumni Admin approves a pending post.", "The post status becomes approved and it may appear in the Graduate feed."),
    ("The Alumni Admin marks a post pending.", "The post status becomes pending and it is not exposed as approved content."),
    ("The Alumni Admin hides a post.", "The post status becomes hidden and it is removed from the approved feed."),
    ("The Alumni Admin reviews a report.", "Reporter, target, reason, and date are shown with the relevant post/comment context."),
    ("The Alumni Admin confirms post deletion.", "The selected post and its comments are removed."),
    ("The Alumni Admin cancels post deletion.", "The post remains."),
    ("The Alumni Admin confirms deletion of one comment.", "Only that comment is removed."),
    ("The Alumni Admin cancels comment deletion.", "The comment remains."),
    ("A moderation update API fails.", "The prior status/content remains and an error is shown."),
    ("A non-Alumni-Admin calls moderation APIs.", "Access is denied."),
])
group("Job Approval Interface", [
    ("The Alumni Admin opens the Pending job tab.", "Pending job posts and summary count are displayed."),
    ("The Alumni Admin filters pending/approved/declined/all.", "Only jobs in the selected status are displayed."),
    ("The Alumni Admin searches by Graduate, program, company, skills, or title.", "Matching jobs are displayed."),
    ("Search produces no job.", "A clear empty state is displayed."),
    ("The Alumni Admin opens job details.", "Poster, requirements, application information, active state, and review data are shown."),
    ("The Alumni Admin enters optional notes and approves a pending job.", "Approval, notes, reviewer, and review date are stored and an active job becomes browseable."),
    ("The Alumni Admin enters optional notes and declines a pending job.", "Declined status and review information are stored and the job remains hidden from Browse Jobs."),
    ("The Alumni Admin cancels an approval/decline confirmation.", "The job remains pending."),
    ("The review API fails.", "The previous approval status remains and an error is displayed."),
    ("An already reviewed item is displayed.", "Prior review notes/time are shown and pending action controls are not incorrectly active."),
    ("A non-Alumni-Admin attempts to review a job.", "Access is denied."),
])
group("Feature Flags and Maintenance Interface", [
    ("Super Admin disables Community Forum.", "Graduate Forum and dependent Messaging navigation/action are unavailable."),
    ("Super Admin disables Messaging while Forum remains enabled.", "Forum remains available and Messaging is unavailable."),
    ("Super Admin disables Job Support.", "Browse Jobs and Job Posting are unavailable."),
    ("Super Admin disables Notifications.", "Notification controls are hidden/unavailable without breaking other pages."),
    ("Super Admin re-enables a feature.", "Its authorized navigation/action becomes available after refreshed configuration."),
    ("Super Admin enables maintenance mode.", "Graduate and non-Super-Admin protected access is blocked."),
    ("Super Admin accesses settings during maintenance.", "Super Admin remains authenticated and able to disable maintenance."),
    ("Super Admin disables maintenance mode.", "Normal role-authorized access resumes."),
])
group("Authorization, Missing Record, and Error Handling Interface", [
    ("A user modifies a URL to request another role's administrator page.", "The route and API deny access without returning protected data."),
    ("A Dean modifies a program parameter outside their allowed scope.", "Out-of-scope graduate/response data is excluded or denied."),
    ("A Graduate modifies an owner ID for a post, comment, profile edit, chat room, or job.", "Ownership and membership checks deny the action."),
    ("A request contains an unknown announcement, survey, response, Graduate, post, room, job, or user ID.", "A controlled not-found/unavailable response is returned."),
    ("A request contains a malformed identifier.", "The request is rejected without an unhandled error or SQL detail."),
    ("The API returns a validation error.", "The interface identifies correctable fields and does not claim success."),
    ("The API returns an authorization error.", "The interface redirects/denies safely and does not expose the response data."),
    ("The API or database becomes unavailable during a read.", "A concise error/empty-safe state and retry where supported are shown."),
    ("The API or database fails during a create/update/delete.", "No duplicate, partial, or false successful operation is presented."),
    ("The same submit control is triggered repeatedly while a request is pending.", "Duplicate records/messages/actions are prevented."),
    ("An authenticated session expires during navigation or form submission.", "The user is asked to sign in again and protected data is not displayed."),
    ("A successful audited administrator action is completed.", "The audit trail records the actor, module, action, and applicable change data without secrets."),
])


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_margins(cell, top=50, start=60, bottom=50, end=60) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "8")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "000000")


def set_table_geometry(table, widths: Sequence[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    set_table_borders(table)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_run_font(run, size: float, bold: bool = False, italic: bool = False, color: str | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def format_table_paragraph(paragraph, center: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.keep_together = False
    paragraph.paragraph_format.widow_control = True


def clear_cell(cell) -> None:
    cell.text = ""
    format_table_paragraph(cell.paragraphs[0])


def write_cell_lines(cell, values: Iterable[str], size: float = 10, bullet: bool = False, bold: bool = False, center: bool = False) -> None:
    clear_cell(cell)
    vals = list(values)
    if not vals:
        return
    for index, value in enumerate(vals):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        format_table_paragraph(paragraph, center=center)
        text_value = f"• {value}" if bullet else value
        run = paragraph.add_run(text_value)
        set_run_font(run, size, bold=bold)


def write_header(cell, value: str, size: float = 10) -> None:
    write_cell_lines(cell, [value], size=size, bold=True, center=True)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, "D9E2F3")


def add_system_table(
    document: Document,
    cases: Sequence[tuple[str, list[str], list[str]]],
    prefix: str,
    id_code: str = "GT",
) -> None:
    headers = ("Test Case Specifications", "Pre-Requisites", "Description", "Expected Result", "Actual Result")
    table = document.add_table(rows=1, cols=5)
    set_table_geometry(table, SYSTEM_WIDTHS)
    header = table.rows[0]
    set_repeat_table_header(header)
    set_cant_split(header)
    header.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for idx, value in enumerate(headers):
        write_header(header.cells[idx], value)

    for number, (description, prereqs, expected) in enumerate(cases, start=1):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        write_cell_lines(row.cells[0], [f"[{prefix}]{id_code}-{number:03d}"], bold=True, center=True)
        write_cell_lines(row.cells[1], prereqs, bullet=True)
        write_cell_lines(row.cells[2], [description])
        write_cell_lines(row.cells[3], expected, bullet=True)
        clear_cell(row.cells[4])
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, SYSTEM_WIDTHS[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_black_box_table(document: Document, cases: Sequence[tuple[str, str]], start_number: int) -> int:
    headers = ("Test ID", "Test Description", "Expected Output", "Actual Output")
    table = document.add_table(rows=1, cols=4)
    set_table_geometry(table, BLACK_BOX_WIDTHS)
    header = table.rows[0]
    set_repeat_table_header(header)
    set_cant_split(header)
    for idx, value in enumerate(headers):
        write_header(header.cells[idx], value)
    current = start_number
    for description, expected in cases:
        row = table.add_row()
        set_cant_split(row)
        write_cell_lines(row.cells[0], [f"BBT-GT-{current:03d}"], bold=True, center=True)
        write_cell_lines(row.cells[1], [description])
        write_cell_lines(row.cells[2], [expected])
        clear_cell(row.cells[3])
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, BLACK_BOX_WIDTHS[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        current += 1
    return current


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    set_run_font(run, 10)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.45)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    for style_name in ("Title", "Heading 1", "Heading 2"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)
    document.styles["Title"].font.size = Pt(14)
    document.styles["Title"].font.bold = True
    document.styles["Heading 1"].font.size = Pt(12)
    document.styles["Heading 1"].font.bold = True
    document.styles["Heading 1"].paragraph_format.space_before = Pt(12)
    document.styles["Heading 1"].paragraph_format.space_after = Pt(6)
    document.styles["Heading 1"].paragraph_format.keep_with_next = True
    document.styles["Heading 2"].font.size = Pt(12)
    document.styles["Heading 2"].font.bold = True
    document.styles["Heading 2"].paragraph_format.space_before = Pt(9)
    document.styles["Heading 2"].paragraph_format.space_after = Pt(4)
    document.styles["Heading 2"].paragraph_format.keep_with_next = True

    add_page_number(section.footer.paragraphs[0])


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def build_document() -> tuple[int, int]:
    document = Document()
    configure_document(document)
    document.core_properties.title = "GradTrack Test Plan"
    document.core_properties.subject = "Complete system interface and black box test plan"
    document.core_properties.author = "GradTrack Project Team"
    document.core_properties.comments = "Generated from the implemented GradTrack routes, interfaces, APIs, roles, and database schema."

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("TEST PLAN")
    set_run_font(title_run, 14, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("GradTrack: A Web-Based Graduate Tracer System with Alumni Job Support System for Norzagaray College")
    set_run_font(run, 12, bold=True)

    scope = document.add_paragraph()
    scope.paragraph_format.space_after = Pt(8)
    scope_run = scope.add_run(
        "Coverage basis: the current React/PHP/MySQL implementation, including Public users, Graduates, Admin, Registrar, "
        "Alumni Admin, Super Admin, Dean-CCS, Dean-COED, Dean-HM, MIS Staff, and Research Coordinator. "
        "Actual Result and Actual Output cells are intentionally blank for manual test execution."
    )
    set_run_font(scope_run, 12, italic=True)

    add_heading(document, "1. System Interface Testing: Web", 1)
    add_system_table(document, web_cases, "WEB")

    add_heading(document, "2. System Interface Testing: Mobile / Responsive Web", 1)
    add_system_table(document, mobile_cases, "MOBILE")

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(document, "3. Black Box Testing", 1)
    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    intro_run = intro.add_run("The following test cases evaluate successful and unsuccessful inputs and observable system behavior. Actual Output is left blank for the tester.")
    set_run_font(intro_run, 12)

    next_id = 1
    for interface, cases in black_box_groups:
        add_heading(document, f"User Interface: {interface}", 2)
        next_id = add_black_box_table(document, cases, next_id)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return len(web_cases) + len(mobile_cases), next_id - 1


def validate_docx(system_count: int, black_count: int) -> None:
    if not OUTPUT.exists() or OUTPUT.suffix.lower() != ".docx":
        raise RuntimeError("DOCX output was not created.")
    if OUTPUT.stat().st_size < 50_000:
        raise RuntimeError("DOCX output is unexpectedly small.")

    reopened = Document(OUTPUT)
    section = reopened.sections[0]
    expected_inches = {
        "page_width": 8.5,
        "page_height": 11.0,
        "left_margin": 1.5,
        "right_margin": 1.0,
        "top_margin": 1.0,
        "bottom_margin": 1.0,
    }
    for name, value in expected_inches.items():
        actual = getattr(section, name).inches
        if abs(actual - value) > 0.01:
            raise RuntimeError(f"Incorrect {name}: {actual}")

    expected_tables = 2 + len(black_box_groups)
    if len(reopened.tables) != expected_tables:
        raise RuntimeError(f"Expected {expected_tables} tables, found {len(reopened.tables)}")

    system_ids: list[str] = []
    black_ids: list[str] = []
    actual_cells_nonblank: list[str] = []
    for table_index, table in enumerate(reopened.tables):
        expected_columns = 5 if table_index < 2 else 4
        expected_widths = SYSTEM_WIDTHS if table_index < 2 else BLACK_BOX_WIDTHS
        if len(table.columns) != expected_columns:
            raise RuntimeError(f"Table {table_index + 1} has an unexpected column count.")
        grid_widths = tuple(int(node.get(qn("w:w"))) for node in table._tbl.tblGrid.gridCol_lst)
        if grid_widths != expected_widths or sum(grid_widths) != USABLE_WIDTH_TWIPS:
            raise RuntimeError(f"Table {table_index + 1} has an incorrect fixed grid: {grid_widths}")
        layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        if layout is None or layout.get(qn("w:type")) != "fixed":
            raise RuntimeError(f"Table {table_index + 1} is not fixed-layout.")
        table_width = table._tbl.tblPr.find(qn("w:tblW"))
        if table_width is None or table_width.get(qn("w:type")) != "dxa" or int(table_width.get(qn("w:w"))) != USABLE_WIDTH_TWIPS:
            raise RuntimeError(f"Table {table_index + 1} does not use the full usable width.")
        if not table.rows[0]._tr.xpath("./w:trPr/w:tblHeader"):
            raise RuntimeError(f"Table {table_index + 1} does not repeat its header row.")
        for row in table.rows:
            cell_widths = tuple(int(cell._tc.tcPr.tcW.get(qn("w:w"))) for cell in row.cells)
            if cell_widths != expected_widths:
                raise RuntimeError(f"Table {table_index + 1} contains a resized row: {cell_widths}")
        for row_index, row in enumerate(table.rows[1:], start=1):
            test_id = row.cells[0].text.strip()
            if table_index < 2:
                system_ids.append(test_id)
                actual_index = 4
            else:
                black_ids.append(test_id)
                actual_index = 3
            if row.cells[actual_index].text.strip():
                actual_cells_nonblank.append(f"table {table_index + 1}, row {row_index + 1}")
    if len(system_ids) != system_count:
        raise RuntimeError(f"System test count mismatch: {len(system_ids)} vs {system_count}")
    if len(black_ids) != black_count:
        raise RuntimeError(f"Black-box count mismatch: {len(black_ids)} vs {black_count}")
    if actual_cells_nonblank:
        raise RuntimeError(f"Actual-result cells are not blank: {actual_cells_nonblank[:3]}")
    if len(system_ids) != len(set(system_ids)) or len(black_ids) != len(set(black_ids)):
        raise RuntimeError("Duplicate test IDs were found.")

    with ZipFile(OUTPUT, "r") as archive:
        names = set(archive.namelist())
        required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
        if not required.issubset(names):
            raise RuntimeError("DOCX package is missing required parts.")
        xml = archive.read("word/document.xml").decode("utf-8")
        if f'w:w="{PAGE_WIDTH_TWIPS}"' not in xml or f'w:h="{PAGE_HEIGHT_TWIPS}"' not in xml:
            raise RuntimeError("Letter page dimensions were not found in document XML.")
        if "w:tblLayout w:type=\"fixed\"" not in xml:
            raise RuntimeError("Fixed table layout was not found.")
        if "w:tblHeader" not in xml:
            raise RuntimeError("Repeating table headers were not found.")
        for width in SYSTEM_WIDTHS + BLACK_BOX_WIDTHS:
            if f'w:gridCol w:w="{width}"' not in xml:
                raise RuntimeError(f"Expected controlled table width {width} was not found.")


if __name__ == "__main__":
    system_total, black_total = build_document()
    validate_docx(system_total, black_total)
    print(f"Created: {OUTPUT}")
    print(f"System Interface Test Cases: {system_total} ({len(web_cases)} web + {len(mobile_cases)} responsive web)")
    print(f"Black Box Test Cases: {black_total}")
    print(f"Black Box Interfaces: {len(black_box_groups)}")
    print(f"File Size: {OUTPUT.stat().st_size} bytes")
