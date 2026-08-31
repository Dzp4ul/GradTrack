from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from generate_gradtrack_test_plan import (
    BLACK_BOX_WIDTHS,
    PAGE_HEIGHT_TWIPS,
    PAGE_WIDTH_TWIPS,
    SYSTEM_WIDTHS,
    USABLE_WIDTH_TWIPS,
    add_black_box_table,
    add_heading,
    add_system_table,
    bullets,
    configure_document,
    set_run_font,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "GradTrack_Test_Plan_Revised.docx"

DB = "Database connection must be established."


web_cases: list[tuple[str, list[str], list[str]]] = []


def web(description: str, prereq: str, expected: str) -> None:
    web_cases.append((description, bullets(prereq), bullets(expected)))


web(
    "Login",
    f"{DB}|User account must exist.",
    "Login interface must load.|Email and password fields must be present.|Required fields must be validated.|Incorrect login details or inactive accounts must show an appropriate message.|Valid login must redirect the user based on the actual account role.",
)
web(
    "Forgot Password",
    f"{DB}|A registered account email must exist.|Email service must be available.",
    "Forgot Password interface must load.|User must be able to request and enter the six-digit verification code.|Incorrect or expired codes must show an appropriate message.|New password and confirmation must be validated.|Password must be updated when valid information is submitted.",
)
web(
    "Graduate Registration",
    f"{DB}|Graduate must have completed the tracer survey.",
    "Registration interface must load after survey submission.|Email and password fields must be validated.|Duplicate Graduate accounts must not be created.|Valid registration must create a pending account.|Graduate must be informed that Alumni Admin approval is required.",
)
web(
    "Graduate Announcements",
    f"{DB}|User must be logged in as a Graduate.|Published announcements must exist.",
    "Announcement list must load.|Graduate must be able to search and filter announcements.|Selected announcement details and images must display properly.|Only published announcements must be shown.|No-result or unavailable records must show an appropriate message.",
)
web(
    "Graduate Dashboard",
    f"{DB}|User must be logged in as a Graduate.",
    "Dashboard must load successfully.|Graduate summary, alumni score, employment information, and available feature tiles must display properly.|Selecting a tile must open the correct Graduate Portal feature.|Unavailable features must be identified clearly.",
)
web(
    "Graduate Profile",
    f"{DB}|User must be logged in as a Graduate.",
    "Graduate profile must load.|User must be able to edit permitted personal information.|Profile and cover photos must be updated properly.|Saved changes must appear after updating.|Password change must work when valid information is provided.",
)
web(
    "Tracer Survey",
    f"{DB}|Active tracer survey must exist.|Graduate record must exist.",
    "Graduate verification and survey interface must load.|Survey questions and sections must display properly.|Required questions and address selections must be validated.|Previous answers must remain while moving between sections.|Completed survey answers must be submitted and saved successfully.",
)
web(
    "Community Forum",
    f"{DB}|User must be logged in as a Graduate.|Community Forum must be enabled.",
    "Forum posts must load.|Graduate must be able to search, filter, create, edit, and delete permitted posts.|Comments and reactions must work.|Supported images or videos must display properly.|Post and comment reporting must be available.",
)
web(
    "Messaging",
    f"{DB}|User must be logged in as a Graduate.|Messaging feature must be available.",
    "Conversation list must load.|Graduate must be able to open direct or group conversations.|Text messages must be sent and received in real time.|Image and supported file attachments must display properly.|Typing indicator and online/offline status must update properly.",
)
web(
    "Alumni Job Support",
    f"{DB}|User must be logged in as a Graduate.|Alumni Job Support must be enabled.",
    "Approved job posts must load.|Graduate must be able to search and view job details and application information.|Eligible employed Graduates must be able to create and edit their own job posts.|New or edited job posts must be submitted for Alumni Admin approval.|Graduate must be able to delete a permitted job post.",
)
web(
    "Notifications",
    f"{DB}|User must be authenticated.|Notifications must be enabled.",
    "Notification list and unread count must load.|User must be able to open a related notification.|One or all notifications must be marked as read.|Read status must remain updated after refreshing.",
)
web(
    "Admin Dashboard",
    f"{DB}|Authorized administrator account must exist.|Survey and Graduate data must exist.",
    "Dashboard must load according to the administrator role.|Survey, graduate, employment, and job-alignment summaries must display properly where authorized.|Available filters must update the displayed information.|Navigation must show only the functions assigned to the role.",
)
web(
    "Graduate Records Management",
    f"{DB}|User must be logged in as Registrar.|Graduate records must be available.",
    "Graduate records must load by program.|Registrar must be able to search and filter records.|Valid Graduate records must be added and edited.|Selected records must be deleted only after confirmation.|Graduate records must be importable from a supported Excel file.",
)
web(
    "Survey Management",
    f"{DB}|User must be logged in with the Admin role.",
    "Survey list must load.|Admin must be able to create and edit survey information, sections, questions, and choices.|Survey preview must show the configured questions.|Only one survey must be active at a time.|Survey deletion must require confirmation.",
)
web(
    "Graduate Participation and Response Monitoring",
    f"{DB}|Authorized Admin or Dean account must exist.|A survey and Graduate records must exist.",
    "Graduate participation list must load.|Search, program, year, and response-status filters must work.|Completed responses must be viewable.|Selected non-respondents must be notified where allowed.|Dean accounts must only see their assigned programs.",
)
web(
    "Reports and Analytics",
    f"{DB}|User must be logged in with the Admin role.|Reportable data must exist.",
    "Reports and charts must load.|Program, year, employment, alignment, salary, location, and survey information must display properly.|Available filters must update report results.|Reports must be exportable to Excel and PDF.|No-data conditions must display properly.",
)
web(
    "User Management",
    f"{DB}|User must be logged in as Super Admin.",
    "Administrator account list must load.|Super Admin must be able to search and filter accounts.|Valid administrator accounts must be created and edited.|Duplicate usernames or emails must be rejected.|Accounts must be activated or deactivated after confirmation.",
)
web(
    "Alumni Verification and Registry",
    f"{DB}|User must be logged in as Alumni Admin.|Graduate accounts or registry records must exist.",
    "Pending Graduate accounts and official alumni records must load.|Alumni Admin must be able to approve or reject a Graduate account.|Registry records must be searchable, editable, verified, linked, or marked inactive.|Alumni lists must be importable from XLSX or CSV.|Registry data must be exportable to CSV or XLSX.",
)
web(
    "Announcement Management",
    f"{DB}|User must be logged in as Alumni Admin.",
    "Announcement list must load.|Alumni Admin must be able to create and edit announcement content.|Cover and gallery images must be uploaded and displayed properly.|Draft, published, and archived statuses must work.|Announcement deletion must require confirmation.",
)
web(
    "Job Post Management",
    f"{DB}|User must be logged in as Alumni Admin.|Graduate job posts must exist.",
    "Pending, approved, and declined job posts must load.|Alumni Admin must be able to search and view job details.|A job post must be approved or declined with optional review notes.|Approved active jobs must appear in the Graduate job list.",
)
web(
    "Community Forum Management",
    f"{DB}|User must be logged in as Alumni Admin.|Forum posts must exist.",
    "Forum moderation list and reports must load.|Posts must be searchable and filterable.|Alumni Admin must be able to approve, mark pending, or hide a post.|Inappropriate posts or comments must be deleted only after confirmation.",
)
web(
    "Auto Email Reminders",
    f"{DB}|User must be logged in as Super Admin.|An active survey must exist.|Email service must be available.",
    "Eligible non-respondents and reminder statistics must load.|Super Admin must be able to send survey reminders after confirmation.|Sent, failed, and skipped results must be displayed.|Reminder interval from 1 to 365 days must be saved.|Reminder history must be viewable.",
)
web(
    "Audit Trail and Database Backup",
    f"{DB}|User must be logged in as Super Admin.",
    "Audit records must load and be searchable or filterable.|Audit details and CSV export must be available.|Database summary and table information must load.|Super Admin must be able to download the database backup as an SQL file.|Unauthorized roles must not access these functions.",
)
web(
    "System Settings",
    f"{DB}|User must be logged in as Super Admin.",
    "General system information, branding, login-page content, and public website content must load and save.|Feature settings for survey, job support, forum, messaging, and notifications must work.|About, FAQ, and Privacy content must be manageable.|Valid images and colors must update the interface.|Maintenance mode must block regular users while keeping Super Admin access.",
)


mobile_cases: list[tuple[str, list[str], list[str]]] = []


def mobile(description: str, prereq: str, expected: str) -> None:
    mobile_cases.append((description, bullets(prereq), bullets(expected)))


mobile(
    "Login",
    f"{DB}|Mobile browser must be available.|User account must exist.",
    "Login interface must fit the mobile screen.|Email, password, show/hide, and submit controls must remain usable.|Validation messages must be readable.|Valid login must open the correct responsive page.|No major element must overlap.",
)
mobile(
    "Graduate Portal",
    f"{DB}|User must be logged in as a Graduate using a mobile browser.",
    "Portal content must fit the screen.|Mobile navigation must remain accessible.|Announcements, dashboard, profile, and notifications must be readable and usable.|Images must resize properly.|No major element must overlap.",
)
mobile(
    "Tracer Survey",
    f"{DB}|Active survey and valid Graduate record must exist.|Mobile browser must be available.",
    "Survey sections and questions must fit the screen.|Choice fields, address fields, and navigation buttons must remain usable.|Required-field messages must be readable.|Entered answers must remain while moving between sections.|Survey must submit successfully.",
)
mobile(
    "Community Forum",
    f"{DB}|Graduate must be logged in on a mobile browser.|Community Forum must be enabled.",
    "Forum posts must fit the screen.|Search, filters, post creation, comments, reactions, and report controls must remain accessible.|Images and videos must resize properly.|Text and buttons must not overlap.",
)
mobile(
    "Messaging",
    f"{DB}|Graduate must be logged in on a mobile browser.|Messaging must be enabled.",
    "Conversation list and chat room must fit the screen.|Message input and send controls must remain visible and usable.|Text, image, and file messages must display properly.|Typing and presence information must remain readable.|Mobile keyboard must not permanently cover the message controls.",
)
mobile(
    "Job Posts",
    f"{DB}|Graduate must be logged in on a mobile browser.|Alumni Job Support must be enabled.",
    "Job list, search, and details must fit the screen.|Application information must remain readable and usable.|Eligible Graduates must be able to complete the job-posting form.|Buttons and fields must remain accessible.|No major element must overlap.",
)


black_box_groups: list[tuple[str, list[tuple[str, str]]]] = [
    ("Login Interface", [
        ("The user enters valid account information.", "Login is successful and the user is redirected according to the account role."),
        ("The user enters incorrect or incomplete login information.", "An appropriate validation or login error message is displayed."),
        ("An inactive, pending, or unauthorized account attempts to log in.", "Access is denied and an appropriate account-status message is displayed."),
    ]),
    ("Forgot Password Interface", [
        ("The user requests and enters a valid verification code, then provides a valid new password.", "The password is updated and the user can return to the correct login page."),
        ("The user enters an invalid or expired code, or mismatched/weak passwords.", "The password is not changed and an appropriate validation message is displayed."),
    ]),
    ("Graduate Registration Interface", [
        ("A survey completer submits a unique email and valid matching password.", "A pending Graduate account is created and approval information is displayed."),
        ("The user submits invalid, incomplete, or duplicate account information.", "Registration is not completed and an appropriate validation message is displayed."),
    ]),
    ("Graduate Portal and Profile Interface", [
        ("The Graduate opens the dashboard, announcements, and profile.", "The correct Graduate information and available portal functions are displayed."),
        ("The Graduate edits permitted profile information and uploads valid profile or cover images.", "The changes are saved and displayed after updating."),
        ("The Graduate submits invalid profile information or unsupported images.", "The changes are not saved and an appropriate message is displayed."),
    ]),
    ("Tracer Survey Interface", [
        ("A verified Graduate completes all required survey information.", "The survey is submitted successfully and the answers are saved."),
        ("The Graduate leaves a required question or address field incomplete.", "The survey is not submitted and the missing information is identified."),
        ("The Graduate moves between survey sections before submitting.", "Entered answers remain displayed in the correct questions."),
    ]),
    ("Community Forum Interface", [
        ("The Graduate creates a valid forum post and comment.", "The post and comment are saved and displayed according to the current moderation status."),
        ("The Graduate reacts to or reports a post/comment.", "The reaction or report is recorded for the selected content."),
        ("The owner edits or confirms deletion of permitted content.", "The selected post or comment is updated or removed; canceled deletion keeps it unchanged."),
    ]),
    ("Messaging Interface", [
        ("The Graduate opens a conversation and sends a text message.", "The text message is sent and received in the selected conversation."),
        ("The Graduate sends a valid image or supported file.", "The attachment is displayed as a message and can be opened by the recipient."),
        ("Two Graduates use the same conversation.", "New messages, typing indicator, and online/offline status update properly."),
    ]),
    ("Alumni Job Support Interface", [
        ("The Graduate searches for and opens an approved job post.", "Matching jobs and complete application information are displayed."),
        ("An eligible Graduate creates or edits a valid job post.", "The job post is saved with pending approval status."),
        ("The Graduate submits incomplete job information or confirms deletion of an owned post.", "Incomplete information is rejected; a confirmed valid deletion removes the selected post."),
    ]),
    ("Graduate Records Interface", [
        ("The Registrar adds or edits a Graduate using valid information.", "The Graduate record is saved and displayed under the correct program and year."),
        ("The Registrar submits incomplete, invalid, or duplicate Graduate information.", "The record is not saved and an appropriate validation message is displayed."),
        ("The Registrar imports a supported Excel file or confirms record deletion.", "Valid rows are imported with a summary, or the selected confirmed records are deleted."),
    ]),
    ("Survey Management and Participation Interface", [
        ("The Admin creates or edits a survey with valid questions and choices.", "The survey is saved and its preview displays the configured content."),
        ("The Admin or Dean searches and filters Graduate participation records.", "Matching answered or non-respondent records are displayed within the user's allowed scope."),
        ("The authorized user views answers or notifies selected non-respondents.", "The correct response is displayed or the reminder result is reported."),
    ]),
    ("Reports and Analytics Interface", [
        ("The Admin selects a report and applies available filters.", "The related summaries, tables, and charts display the filtered information."),
        ("The Admin exports the selected report.", "A readable Excel or PDF report is downloaded using the selected scope."),
    ]),
    ("User Management Interface", [
        ("The Super Admin creates or edits an administrator using valid information.", "The administrator account is saved with the selected role and status."),
        ("The Super Admin submits duplicate or incomplete account information.", "The account is not saved and an appropriate validation message is displayed."),
        ("The Super Admin confirms account activation or deactivation.", "The selected account status is updated; the logged-in Super Admin cannot deactivate their own account."),
    ]),
    ("Alumni Verification and Registry Interface", [
        ("The Alumni Admin approves or rejects a pending Graduate account.", "The selected verification status and optional reason are saved."),
        ("The Alumni Admin searches, edits, verifies, links, or marks an alumni registry record inactive.", "The selected official alumni record is updated correctly."),
        ("The Alumni Admin imports or exports an alumni list.", "Valid XLSX/CSV records are processed with a summary, or the selected registry scope is downloaded."),
    ]),
    ("Announcement Management Interface", [
        ("The Alumni Admin creates or edits an announcement with valid content and images.", "The announcement is saved with its selected status and media."),
        ("The Alumni Admin publishes an announcement.", "The announcement becomes visible in the Graduate Portal."),
        ("The Alumni Admin confirms or cancels announcement deletion.", "Confirmed deletion removes the selected announcement; Cancel keeps it unchanged."),
    ]),
    ("Forum and Job Post Management Interface", [
        ("The Alumni Admin approves, hides, or marks a forum post pending.", "The selected forum status is updated and Graduate visibility follows that status."),
        ("The Alumni Admin approves or declines a pending job post.", "The decision and optional review notes are saved; approved active jobs become visible."),
        ("The Alumni Admin confirms deletion of inappropriate forum content.", "The selected post or comment is removed; canceled deletion keeps it unchanged."),
    ]),
    ("System Administration Interface", [
        ("The Super Admin updates valid system, branding, feature, or public website settings.", "The settings are saved and displayed in the related interfaces."),
        ("The Super Admin sends survey reminders or changes the reminder interval.", "Eligible graduates are processed, results are displayed, and the valid interval is saved."),
        ("The Super Admin views/exports audit records or downloads a database backup.", "Audit information or the SQL backup file is produced successfully."),
        ("The Super Admin enables maintenance mode.", "Regular protected users are blocked while Super Admin remains able to restore access."),
    ]),
]


def build_document() -> tuple[int, int]:
    document = Document()
    configure_document(document)
    document.core_properties.title = "GradTrack Revised Test Plan"
    document.core_properties.subject = "Simple academic system interface and black box test plan"
    document.core_properties.author = "GradTrack Project Team"

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("TEST PLAN"), 14, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    set_run_font(
        subtitle.add_run("GradTrack: A Web-Based Graduate Tracer System with Alumni Job Support System for Norzagaray College"),
        12,
        bold=True,
    )

    add_heading(document, "1. System Interface Testing: Web", 1)
    add_system_table(document, web_cases, "WEB", id_code="GTS")

    add_heading(document, "2. System Interface Testing: Mobile", 1)
    add_system_table(document, mobile_cases, "MOBILE", id_code="GTS")

    add_heading(document, "3. Black Box Testing", 1)
    next_id = 1
    for interface, cases in black_box_groups:
        add_heading(document, f"User Interface: {interface}", 2)
        next_id = add_black_box_table(document, cases, next_id)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return len(web_cases) + len(mobile_cases), next_id - 1


def validate_document(system_count: int, black_count: int) -> None:
    if not OUTPUT.exists() or OUTPUT.suffix.lower() != ".docx":
        raise RuntimeError("Revised DOCX was not created.")
    with ZipFile(OUTPUT, "r") as archive:
        if archive.testzip() is not None:
            raise RuntimeError("The DOCX package failed its ZIP integrity check.")
        xml = archive.read("word/document.xml").decode("utf-8")
        if f'w:w="{PAGE_WIDTH_TWIPS}"' not in xml or f'w:h="{PAGE_HEIGHT_TWIPS}"' not in xml:
            raise RuntimeError("Letter page dimensions were not found.")
        if "Online STEP" in xml:
            raise RuntimeError("Unrelated reference-system content was found.")

    reopened = Document(OUTPUT)
    section = reopened.sections[0]
    dimensions = (
        section.page_width.inches,
        section.page_height.inches,
        section.left_margin.inches,
        section.right_margin.inches,
        section.top_margin.inches,
        section.bottom_margin.inches,
    )
    expected_dimensions = (8.5, 11.0, 1.5, 1.0, 1.0, 1.0)
    if any(abs(actual - expected) > 0.01 for actual, expected in zip(dimensions, expected_dimensions)):
        raise RuntimeError(f"Incorrect page setup: {dimensions}")

    if len(reopened.tables) != 2 + len(black_box_groups):
        raise RuntimeError("Unexpected table count.")

    ids: list[str] = []
    for table_index, table in enumerate(reopened.tables):
        expected_widths = SYSTEM_WIDTHS if table_index < 2 else BLACK_BOX_WIDTHS
        expected_columns = len(expected_widths)
        if len(table.columns) != expected_columns:
            raise RuntimeError(f"Table {table_index + 1} has the wrong column count.")
        grid = tuple(int(node.get(qn("w:w"))) for node in table._tbl.tblGrid.gridCol_lst)
        if grid != expected_widths or sum(grid) != USABLE_WIDTH_TWIPS:
            raise RuntimeError(f"Table {table_index + 1} has inconsistent widths: {grid}")
        layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        if layout is None or layout.get(qn("w:type")) != "fixed":
            raise RuntimeError(f"Table {table_index + 1} is not fixed width.")
        if not table.rows[0]._tr.xpath("./w:trPr/w:tblHeader"):
            raise RuntimeError(f"Table {table_index + 1} does not repeat its header.")
        actual_column = 4 if table_index < 2 else 3
        for row in table.rows[1:]:
            ids.append(row.cells[0].text.strip())
            if row.cells[actual_column].text.strip():
                raise RuntimeError(f"Result cell is not blank in {row.cells[0].text.strip()}.")
            widths = tuple(int(cell._tc.tcPr.tcW.get(qn("w:w"))) for cell in row.cells)
            if widths != expected_widths:
                raise RuntimeError(f"A row in table {table_index + 1} was resized.")

    if len(ids) != system_count + black_count or len(ids) != len(set(ids)):
        raise RuntimeError("Test ID count or uniqueness validation failed.")
    if len(web_cases) > 25 or not 5 <= len(mobile_cases) <= 8:
        raise RuntimeError("The revised interface-test count exceeds the requested scope.")
    if any(not 1 <= len(cases) <= 4 for _, cases in black_box_groups):
        raise RuntimeError("A black-box interface has too many cases.")


if __name__ == "__main__":
    system_total, black_total = build_document()
    validate_document(system_total, black_total)
    print(f"Created: {OUTPUT}")
    print(f"System Interface Test Cases: {system_total} ({len(web_cases)} web + {len(mobile_cases)} mobile)")
    print(f"Black Box Test Cases: {black_total}")
    print(f"Black Box Interfaces: {len(black_box_groups)}")
    print(f"File Size: {OUTPUT.stat().st_size} bytes")
